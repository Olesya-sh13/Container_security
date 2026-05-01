
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
sec.left_margin = Cm(3.0);  sec.right_margin = Cm(1.5)
sec.top_margin  = Cm(2.0);  sec.bottom_margin = Cm(2.0)

TNR = "Times New Roman"
CNR = "Courier New"

def fmt(para, size=14, bold=False, italic=False,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        sb=0, sa=6, fi=None, font=TNR):
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if fi is not None: pf.first_line_indent = Cm(fi)
    for run in para.runs:
        run.font.name = font; run.font.size = Pt(size)
        run.font.bold = bold; run.font.italic = italic

def h1(text):
    p = doc.add_paragraph(); p.add_run(text)
    fmt(p, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=12, sa=6, fi=0)

def h2(text):
    p = doc.add_paragraph(); p.add_run(text)
    fmt(p, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=4, fi=0)

def h3(text):
    p = doc.add_paragraph(); p.add_run(text)
    fmt(p, size=14, bold=False, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, sb=4, sa=4, fi=0)

def body(text):
    p = doc.add_paragraph(); p.add_run(text)
    fmt(p, fi=1.25)

def bul(text):
    p = doc.add_paragraph(style='List Bullet'); p.add_run(text)
    fmt(p, fi=0)

def code(text):
    p = doc.add_paragraph()
    p.add_run(text)
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0); pf.left_indent = Cm(1.0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs:
        run.font.name = CNR; run.font.size = Pt(10)

def pb(): doc.add_page_break()

def table_caption(text):
    p = doc.add_paragraph(); p.add_run(text)
    fmt(p, size=12, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, fi=0, sa=3)

def mk_table(header_row, data_rows, font_size=11):
    tbl = doc.add_table(rows=1+len(data_rows), cols=len(header_row))
    tbl.style = 'Table Grid'
    for j, h in enumerate(header_row):
        tbl.rows[0].cells[j].text = h
    for i, row in enumerate(data_rows, start=1):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = val
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = TNR; run.font.size = Pt(font_size)
    return tbl

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
def tp(txt, **kw): p = doc.add_paragraph(); p.add_run(txt); fmt(p, **kw)

tp("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ",
   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0, sa=2)
tp("Кафедра программной инженерии",
   italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0, sa=50)
tp("ОТЧЁТ ПО ЛАБОРАТОРНОЙ РАБОТЕ",
   bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0, sa=6)
tp("по дисциплине «Сопровождение программных средств»",
   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0, sa=6)
tp("Задания 4–7: Разработка модифицированного контейнера (HTTPS),\n"
   "тестирование по ГОСТ Р 56920 и IEEE 829,\n"
   "размещение у заказчика, оповещение сторон",
   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0, sa=48)
p = doc.add_paragraph(); p.add_run(
    "Выполнил: студент группы ____\n"
    "_______________________\n«___» ____________ 2026 г.")
fmt(p, fi=0, sa=6)
p = doc.add_paragraph(); p.add_run(
    "Проверил: ____________________\n"
    "_______________________\n«___» ____________ 2026 г.")
fmt(p, fi=0, sa=60)
tp("Москва – 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fi=0)
pb()

# ════════════════════════════════════════════════════════════════════════════
# CONTENTS
# ════════════════════════════════════════════════════════════════════════════
h1("СОДЕРЖАНИЕ")
toc = [
    "Введение .............................................................. 4",
    "1  Задание 4. Разработка модифицированного программного продукта (контейнера) .... 5",
    "   1.1  Алгоритм действий при выполнении задания .................. 5",
    "   1.2  Алгоритм работы контейнера КЖД v2.0 ....................... 7",
    "   1.3  Блок-схема алгоритма по ГОСТ 19.701 ........................ 8",
    "   1.4  Фрагменты программного кода с комментариями ................ 9",
    "2  Задание 5. Тестирование по ГОСТ Р 56920 и IEEE 829 ............. 13",
    "   2.1  Описание тестирования ...................................... 13",
    "   2.2  Мастер-план тестирования (IEEE 829, Master Test Plan) ..... 14",
    "   2.3  Спецификация процедуры тестирования (IEEE 829, STP) ........ 17",
    "   2.4  Спецификация тестов (Test Specification) .................. 19",
    "   2.5  Результаты выполнения тестов ............................... 22",
    "3  Задание 6. Размещение обновлённых продуктов в среде заказчика . 23",
    "4  Задание 7. Доведение сведений о модификации до заинтересованных сторон .... 26",
    "Заключение .......................................................... 28",
    "Список использованных источников ................................... 29",
]
for line in toc:
    p = doc.add_paragraph(); p.add_run(line)
    fmt(p, fi=0, sa=2)
pb()

# ════════════════════════════════════════════════════════════════════════════
# ВВЕДЕНИЕ
# ════════════════════════════════════════════════════════════════════════════
h1("ВВЕДЕНИЕ")
body("Настоящий документ является продолжением отчёта по заданиям 1–2 и содержит "
     "материалы по заданиям 4–7 лабораторной работы по дисциплине «Сопровождение "
     "программных средств». Документ разработан в соответствии с требованиями "
     "ГОСТ Р ИСО/МЭК 12207–2010, ГОСТ Р 56920–2016 (ISO/IEC/IEEE 29119), "
     "IEEE 829-2008 и ГОСТ 19.701–90.")
body("Объект разработки — контейнер журналирования событий доступа (КЖД) версии 2.0, "
     "модифицированный для поддержки протокола HTTPS. Приложение развёртывается "
     "в SaaS-среде Replit в рамках задания 4.")
body("Структура документа:")
bul("Раздел 1 (Задание 4): алгоритм действий, алгоритм работы КЖД, блок-схема "
    "по ГОСТ 19.701, фрагменты кода.")
bul("Раздел 2 (Задание 5): план тестирования, спецификация процедуры тестирования, "
    "спецификация тестов по ГОСТ Р 56920 и IEEE 829.")
bul("Раздел 3 (Задание 6): описание процесса размещения обновлённых продуктов "
    "в среде заказчика.")
bul("Раздел 4 (Задание 7): порядок доведения сведений о модификации до "
    "затронутых сторон.")
pb()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 1 — ЗАДАНИЕ 4
# ════════════════════════════════════════════════════════════════════════════
h1("1  ЗАДАНИЕ 4. РАЗРАБОТКА МОДИФИЦИРОВАННОГО\nПРОГРАММНОГО ПРОДУКТА (КОНТЕЙНЕРА)")

h2("1.1  Алгоритм действий при выполнении задания")
body("Ниже приведён пошаговый алгоритм действий, выполненных при реализации "
     "задания 4 в SaaS-среде Replit.")

h3("Шаг 1. Анализ существующего кода (ЛР1)")
body("Выполнено изучение исходного кода приложения (Django, Python 3.10), "
     "структуры логов (audit.log, *.enc), модели пользователей и механизма "
     "шифрования Fernet. Определены точки изменения для HTTPS-перехода.")

h3("Шаг 2. Проектирование целевой архитектуры")
body("Спроектирована трёхконтейнерная архитектура: Nginx (TLS-терминация) + "
     "app (Django) + logger (КЖД). Выбран Nginx 1.26 LTS. Определены: формат "
     "журнала extended_ssl, механизмы доставки (file tail + syslog UDP).")

h3("Шаг 3. Создание конфигурации Nginx с HTTPS")
body("Разработан файл nginx/nginx.conf с директивами: ssl_protocols TLSv1.2 "
     "TLSv1.3, наборы шифров Mozilla Intermediate, HSTS (max-age=31536000), "
     "перенаправление HTTP→HTTPS (301), расширенный лог-формат extended_ssl.")

h3("Шаг 4. Разработка Dockerfile (multi-stage)")
body("Создан многоэтапный Dockerfile: этап builder (python:3.12-alpine3.21) "
     "устанавливает зависимости через Poetry; этап runtime копирует только "
     "виртуальное окружение и код, запускает от непривилегированного пользователя "
     "appuser. Добавлена директива HEALTHCHECK.")

h3("Шаг 5. Разработка агента КЖД v2.0 (log_agent.py)")
body("Реализован агент на Python с двумя параллельными потоками: FileWatcher "
     "(tail метода чтения файла) и SyslogServer (UDP 514). Добавлено регулярное "
     "выражение для парсинга формата extended_ssl (поля ssl_proto, ssl_cipher, "
     "request_time). Реализован REST API (Flask + waitress) на порту 8080.")

h3("Шаг 6. Создание docker-compose.yml")
body("Разработан docker-compose.yml v3.9 с тремя сервисами, именованными томами "
     "(app_data, nginx_logs, logger_data), двумя сетями (frontend — внешняя, "
     "internal — изолированная). Сервис app не публикует порты наружу.")

h3("Шаг 7. Разработка тестов")
body("Разработан файл tests/test_https_container.py с 41 тестом, охватывающим "
     "парсинг (TC-PARSE), базу данных (TC-DB), REST API (TC-API), безопасность "
     "конфигурации (TC-SEC) и интеграцию (TC-INT). Все 41 тест проходят успешно.")

h3("Шаг 8. Генерация самоподписанного сертификата")
body("Создан скрипт certs/generate_self_signed.sh для автоматической генерации "
     "X.509 v3 сертификата с SAN для использования в тестовом окружении.")

h2("1.2  Алгоритм работы контейнера КЖД v2.0")
body("Алгоритм работы КЖД v2.0 состоит из следующих последовательных и параллельных "
     "этапов (см. также блок-схему в разделе 1.3):")

body("1. ИНИЦИАЛИЗАЦИЯ.")
bul("Считываются переменные окружения: LOG_FORMAT, LOG_LEVEL, NGINX_LOG_PATH, "
    "SYSLOG_PORT, DB_PATH, API_PORT.")
bul("Настраивается подсистема логирования агента (logging.basicConfig).")
bul("Выполняется инициализация SQLite-базы данных: создаётся таблица access_events "
    "(если отсутствует) и индекс idx_recorded_at для ускорения запросов.")

body("2. ЗАПУСК ПАРАЛЛЕЛЬНЫХ ПОТОКОВ.")
bul("Поток A — FileWatcher: ожидает появления файла NGINX_LOG_PATH, затем "
    "открывает его, перемещает позицию в конец (seek to end) и в цикле читает "
    "новые строки с паузой 0.5 с при отсутствии данных.")
bul("Поток B — SyslogServer: создаёт UDP-сокет, привязывается к порту SYSLOG_PORT "
    "(514), входит в цикл приёма датаграмм.")

body("3. ОБРАБОТКА СОБЫТИЯ (выполняется в каждом потоке).")
bul("Получение строки (из файла или UDP-датаграммы).")
bul("Предобработка: если строка содержит RFC-3164 syslog-префикс — префикс отбрасывается.")
bul("Парсинг строки регулярным выражением LOG_RE формата extended_ssl.")
bul("При успешном парсинге: формирование словаря события с полями remote_addr, "
    "time_local, method, uri, status, ssl_proto, ssl_cipher, request_time, source.")
bul("Сохранение события в SQLite (функция save_event).")
bul("При неудачном парсинге: запись предупреждения в лог агента, строка пропускается.")

body("4. REST API (выполняется в главном потоке).")
bul("GET /api/v1/health — возврат статуса работоспособности.")
bul("GET /api/v1/logs — выборка событий с фильтрацией по протоколу, датам, "
    "с поддержкой limit/offset.")
bul("GET /api/v1/stats/tls — агрегированная статистика по TLS-протоколам и шифрам.")
bul("POST /api/v1/rotate — запрос ротации журнала.")
bul("API запущен через production WSGI-сервер waitress.")

h2("1.3  Блок-схема алгоритма работы КЖД v2.0 (по ГОСТ 19.701)")
body("Ниже представлено текстовое описание блок-схемы. Символы соответствуют "
     "ГОСТ 19.701–90 (терминатор, процесс, решение, ввод-вывод, соединитель).")

mk_table(
    ["№", "Символ (ГОСТ 19.701)", "Обозначение", "Описание действия"],
    [
        ("1",  "Терминатор",  "НАЧАЛО",            "Запуск контейнера КЖД"),
        ("2",  "Процесс",     "INIT_ENV",           "Чтение переменных окружения"),
        ("3",  "Процесс",     "INIT_DB",            "Инициализация SQLite (CREATE TABLE IF NOT EXISTS)"),
        ("4",  "Процесс",     "START_THREADS",      "Запуск потоков FileWatcher и SyslogServer"),
        ("5",  "Процесс",     "START_API",          "Запуск REST API (waitress)"),
        ("6",  "Ввод/вывод",  "GET_LINE",           "Получение строки (файл / UDP)"),
        ("7",  "Решение",     "HAS_SYSLOG_PREFIX?", "Строка содержит syslog-префикс?"),
        ("8",  "Процесс",     "STRIP_PREFIX",       "Удаление syslog-префикса (Да→8, Нет→9)"),
        ("9",  "Решение",     "MATCH_RE?",          "Строка соответствует LOG_RE?"),
        ("10", "Процесс",     "LOG_WARN",           "Запись предупреждения (Нет→6)"),
        ("11", "Процесс",     "BUILD_EVENT",        "Формирование словаря события (Да)"),
        ("12", "Ввод/вывод",  "SAVE_EVENT",         "Запись события в SQLite"),
        ("13", "Соединитель", "→6",                 "Возврат к ожиданию следующей строки"),
        ("14", "Терминатор",  "КОНЕЦ",              "Остановка контейнера (SIGTERM)"),
    ]
)
table_caption("Таблица 1.1 — Описание блок-схемы алгоритма КЖД v2.0 (ГОСТ 19.701)")

h2("1.4  Фрагменты программного кода с комментариями")

h3("1.4.1  Функция parse_line (docker/logger/log_agent.py)")
body("Функция parse_line выполняет разбор строки журнала Nginx формата extended_ssl. "
     "Ключевые особенности: поддержка syslog-префикса RFC-3164, типизация полей "
     "(status → int, request_time → float), возврат None при неудачном парсинге.")

code("# Регулярное выражение для парсинга строк формата extended_ssl")
code("LOG_RE = re.compile(")
code("    r'(?P<remote_addr>\\S+)'")
code("    r'\\s+\\[(?P<time_local>[^\\]]+)\\]'")
code("    r'\\s+\"(?P<request>[^\"]*)'")
code("    r'\\s+(?P<status>\\d{3})'")
code("    r'\\s+(?P<bytes_sent>\\d+)'")
code("    r'\\s+\"(?P<referer>[^\"]*)'")
code("    r'\\s+\"(?P<user_agent>[^\"]*)'")
code("    r'\\s+ssl_proto=(?P<ssl_proto>\\S+)'")
code("    r'\\s+ssl_cipher=(?P<ssl_cipher>\\S+)'")
code("    r'\\s+rt=(?P<request_time>\\S+)'")
code(")")
code("")
code("def parse_line(line: str, source: str = 'file') -> Optional[dict]:")
code("    line = line.strip()")
code("    if not line:")
code("        return None  # Пустая строка — пропускаем")
code("")
code("    # Удаление syslog-префикса RFC-3164, если присутствует")
code("    syslog_prefix = re.match(r'^<\\d+>.*?nginx:\\s*', line)")
code("    if syslog_prefix:")
code("        line = line[syslog_prefix.end():]")
code("")
code("    m = LOG_RE.match(line)")
code("    if not m:")
code("        logger.debug('Не удалось разобрать строку: %s', line[:120])")
code("        return None")
code("")
code("    req_parts    = m.group('request').split()")
code("    method       = req_parts[0] if len(req_parts) > 0 else '-'")
code("    uri          = req_parts[1] if len(req_parts) > 1 else '-'")
code("    http_version = req_parts[2] if len(req_parts) > 2 else '-'")
code("")
code("    return {")
code("        'recorded_at':  datetime.utcnow().isoformat(),")
code("        'remote_addr':  m.group('remote_addr'),")
code("        'ssl_proto':    m.group('ssl_proto'),   # Версия TLS")
code("        'ssl_cipher':   m.group('ssl_cipher'),  # Набор шифров")
code("        'request_time': float(m.group('request_time')),")
code("        'source':       source,  # 'file' или 'syslog'")
code("        # ... остальные поля ...")
code("    }")

doc.add_paragraph()
h3("1.4.2  Функция file_watcher (docker/logger/log_agent.py)")
body("Реализует tail-чтение файла журнала Nginx. Перемещение позиции в конец файла "
     "при старте предотвращает повторную обработку исторических записей.")

code("def file_watcher() -> None:")
code("    # Ожидание появления файла (Nginx запускается позже КЖД)")
code("    while not os.path.exists(NGINX_LOG):")
code("        time.sleep(2)")
code("")
code("    with open(NGINX_LOG, 'r', encoding='utf-8', errors='replace') as fh:")
code("        fh.seek(0, 2)  # Переход в конец файла (tail-режим)")
code("        while True:")
code("            line = fh.readline()")
code("            if line:")
code("                event = parse_line(line, source='file')")
code("                if event:")
code("                    save_event(event)  # Сохранение в SQLite")
code("            else:")
code("                time.sleep(0.5)  # Пауза при отсутствии новых данных")

doc.add_paragraph()
h3("1.4.3  REST API-эндпоинт /api/v1/logs (docker/logger/log_agent.py)")
body("Обеспечивает фильтрацию событий по версии TLS-протокола, временному диапазону, "
     "с поддержкой пагинации через параметры limit и offset.")

code("@app_api.route('/api/v1/logs', methods=['GET'])")
code("def get_logs():")
code("    protocol = request.args.get('protocol')  # Фильтр по TLS-версии")
code("    from_dt  = request.args.get('from')      # ISO 8601 начало периода")
code("    to_dt    = request.args.get('to')         # ISO 8601 конец периода")
code("    limit    = int(request.args.get('limit',  100))")
code("    offset   = int(request.args.get('offset', 0))")
code("")
code("    query  = 'SELECT * FROM access_events WHERE 1=1'")
code("    params = []")
code("    if protocol:")
code("        query  += ' AND ssl_proto = ?'")
code("        params.append(protocol)")
code("    if from_dt:")
code("        query  += ' AND recorded_at >= ?'")
code("        params.append(from_dt)")
code("    query += ' ORDER BY recorded_at DESC LIMIT ? OFFSET ?'")
code("    params.extend([limit, offset])")
code("")
code("    conn = sqlite3.connect(DB_PATH)")
code("    conn.row_factory = sqlite3.Row")
code("    rows = conn.execute(query, params).fetchall()")
code("    conn.close()")
code("    return jsonify([dict(r) for r in rows])")

doc.add_paragraph()
h3("1.4.4  Конфигурация TLS в nginx/nginx.conf")
body("Ключевые директивы безопасности HTTPS-конфигурации Nginx:")

code("# Разрешены только актуальные версии TLS")
code("ssl_protocols TLSv1.2 TLSv1.3;")
code("")
code("# Наборы шифров Mozilla Intermediate (без RC4, 3DES, экспортных)")
code("ssl_ciphers ECDHE-ECDSA-AES128-GCM-SHA256:ECDHE-RSA-AES128-GCM-SHA256:...")
code("")
code("# Отключение SSL session tickets (обеспечение forward secrecy)")
code("ssl_session_tickets off;")
code("")
code("# Заголовок HSTS: принудительный HTTPS на 1 год")
code("add_header Strict-Transport-Security 'max-age=31536000; includeSubDomains' always;")
code("")
code("# Перенаправление HTTP → HTTPS")
code("server {")
code("    listen 80;")
code("    return 301 https://$host$request_uri;")
code("}")
pb()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 2 — ЗАДАНИЕ 5
# ════════════════════════════════════════════════════════════════════════════
h1("2  ЗАДАНИЕ 5. ТЕСТИРОВАНИЕ ПО ГОСТ Р 56920 И IEEE 829")

h2("2.1  Описание тестирования")
body("Тестирование КЖД v2.0 выполнено в соответствии с:")
bul("ГОСТ Р 56920–2016 (ISO/IEC/IEEE 29119-1:2013) — понятия и определения "
    "тестирования программного обеспечения;")
bul("IEEE 829-2008 — стандарт документирования тестирования программных систем.")
body("Применяемые виды тестирования (по ГОСТ Р 56920, раздел 5.3):")
bul("Модульное (unit) тестирование — отдельные функции parse_line, save_event, init_db.")
bul("Функциональное (functional) тестирование — REST API эндпоинты.")
bul("Тестирование безопасности — конфигурационные файлы (nginx.conf, Dockerfile, "
    "docker-compose.yml).")
bul("Интеграционное тестирование — взаимодействие FileWatcher + SQLite.")

mk_table(
    ["Характеристика", "Значение"],
    [
        ("Всего тест-кейсов",     "41"),
        ("Тест-кейсов TC-PARSE",  "15 (парсинг строк журнала)"),
        ("Тест-кейсов TC-DB",     "5 (операции с БД)"),
        ("Тест-кейсов TC-API",    "10 (REST API)"),
        ("Тест-кейсов TC-SEC",    "10 (безопасность конфигурации)"),
        ("Тест-кейсов TC-INT",    "1 (интеграция)"),
        ("Инструмент",            "pytest 9.0.3 + Python 3.10"),
        ("Статус",                "41/41 PASSED"),
    ]
)
table_caption("Таблица 2.1 — Сводная характеристика тестового покрытия")

h2("2.2  Мастер-план тестирования (IEEE 829, Master Test Plan)")
body("Идентификатор документа: MTP-KJD-2026-001")
body("Версия: 1.0   Дата: 01.05.2026   Автор: Компания 2")

h3("2.2.1  Введение и область применения")
body("Настоящий мастер-план тестирования охватывает верификацию и валидацию "
     "контейнера журналирования событий доступа (КЖД) версии 2.0 на предмет "
     "соответствия функциональным требованиям и требованиям безопасности, "
     "сформулированным в запросе на модификацию ЗМ-001 (переход HTTP → HTTPS).")

h3("2.2.2  Ссылочные документы")
bul("ЗМ-001 — Запрос на модификацию. Переход на HTTPS.")
bul("ГОСТ Р 56920–2016 (ISO/IEC/IEEE 29119-1:2013).")
bul("IEEE 829-2008 — Standard for Software and System Test Documentation.")
bul("ГОСТ Р ИСО/МЭК 12207–2010, раздел 6.4.7 (процесс сопровождения).")

h3("2.2.3  Цели тестирования")
bul("T-OBJ-01: убедиться, что parse_line корректно разбирает все поля "
    "формата extended_ssl, включая ssl_proto и ssl_cipher.")
bul("T-OBJ-02: убедиться, что события сохраняются в SQLite без потери данных.")
bul("T-OBJ-03: убедиться, что REST API возвращает корректные данные и "
    "поддерживает фильтрацию.")
bul("T-OBJ-04: убедиться, что конфигурационные файлы не содержат небезопасных "
    "настроек TLS (TLS 1.0, TLS 1.1, RC4, export ciphers).")
bul("T-OBJ-05: убедиться, что приложение не запускается от пользователя root.")

h3("2.2.4  Критерии приостановки и возобновления тестирования")
body("По ГОСТ Р 56920, раздел 5.5.3 (Suspension and resumption criteria):")
bul("Приостановка: обнаружение критической ошибки в инициализации БД, "
    "не позволяющей выполнить более 50% тестов.")
bul("Возобновление: устранение критической ошибки, подтверждённое успешным "
    "прохождением тестов TC-DB-001 и TC-DB-002.")

h3("2.2.5  Критерии завершения тестирования")
bul("Пройдено не менее 95% тест-кейсов (39/41).")
bul("Отсутствуют открытые дефекты с приоритетом CRITICAL или HIGH.")
bul("Все тесты из группы TC-SEC пройдены (100%).")

h3("2.2.6  Подход к тестированию")
body("Тестирование выполняется в рамках CI-пайплайна командой: "
     "python -m pytest tests/test_https_container.py -v --tb=short. "
     "Тесты TC-SEC осуществляют статический анализ конфигурационных файлов "
     "без запуска Docker-контейнеров, что позволяет выполнять их в среде Replit.")

h3("2.2.7  Риски и допущения")
mk_table(
    ["Риск", "Вероятность", "Влияние", "Меры снижения"],
    [
        ("Docker недоступен в Replit", "Высокая", "Контейнерные тесты невозможны",
         "Тесты TC-SEC и TC-PARSE выполнены статически без Docker"),
        ("Несовместимость Python 3.10 / 3.12", "Низкая", "Сбой тестов",
         "Тесты выполнены на Python 3.10, код совместим с 3.12"),
        ("Отсутствие реального TLS-сертификата", "Высокая", "TC-SEC работают",
         "TC-SEC проверяют конфиги статически, не требуют живого TLS"),
    ]
)
table_caption("Таблица 2.2 — Риски тестирования и меры снижения")

h2("2.3  Спецификация процедуры тестирования (IEEE 829, Test Procedure Specification)")
body("Идентификатор: TPS-KJD-2026-001   Версия: 1.0")
body("Ссылка на план: MTP-KJD-2026-001")

h3("2.3.1  Порядок выполнения тестирования")
body("Процедура выполняется в следующей последовательности:")
bul("Шаг 1. Подготовка среды: установка зависимостей командой "
    "pip install flask waitress pyyaml apachelogs pytest.")
bul("Шаг 2. Запуск тестов группы TC-DB: python -m pytest tests/test_https_container.py "
    "::TestDatabase -v. Убедиться, что все 5 тестов PASSED.")
bul("Шаг 3. Запуск тестов группы TC-PARSE: python -m pytest tests/test_https_container.py "
    "::TestParseLineExtendedSSL -v. Убедиться, что все 15 тестов PASSED.")
bul("Шаг 4. Запуск тестов группы TC-API: python -m pytest tests/test_https_container.py "
    "::TestRestAPI -v. Убедиться, что все 10 тестов PASSED.")
bul("Шаг 5. Запуск тестов группы TC-SEC: python -m pytest tests/test_https_container.py "
    "::TestHTTPSSecurityRequirements -v. Убедиться, что все 10 тестов PASSED.")
bul("Шаг 6. Запуск тестов группы TC-INT: python -m pytest tests/test_https_container.py "
    "::TestIntegrationFileTailing -v. Убедиться, что 1 тест PASSED.")
bul("Шаг 7. Запуск полного набора: python -m pytest tests/test_https_container.py -v "
    "--tb=short. Сверить итог: 41 passed.")
bul("Шаг 8. Фиксация результатов в протоколе тестирования (Test Log, IEEE 829).")

h3("2.3.2  Требования к среде выполнения тестов")
mk_table(
    ["Компонент", "Требование"],
    [
        ("ОС",              "Linux (Replit / Ubuntu 22.04)"),
        ("Python",          "3.10+ (тестировалось на 3.10.19)"),
        ("pytest",          "9.0.3+"),
        ("Flask",           "3.1.0+"),
        ("waitress",        "3.0.1+"),
        ("PyYAML",          "6.0+"),
        ("Файлы",           "nginx/nginx.conf, Dockerfile, docker-compose.yml должны существовать"),
    ]
)
table_caption("Таблица 2.3 — Требования к среде тестирования")

h2("2.4  Спецификация тестов (Test Specification)")
body("Ниже приведена сводная таблица тест-кейсов с идентификаторами, входными данными, "
     "ожидаемыми результатами и критериями прохождения (по IEEE 829, Test Case Specification).")

mk_table(
    ["ID", "Класс", "Наименование", "Входные данные", "Ожидаемый результат", "Статус"],
    [
        ("TC-PARSE-001","PARSE","TLS1.3: возврат dict",
         "Строка TLS1.3 extended_ssl","dict, not None","PASS"),
        ("TC-PARSE-002","PARSE","Корректный IP-адрес",
         "Строка TLS1.3","remote_addr == '192.168.1.100'","PASS"),
        ("TC-PARSE-003","PARSE","Метод и URI",
         "Строка TLS1.3","method='GET', uri='/logger/'","PASS"),
        ("TC-PARSE-004","PARSE","HTTP-статус",
         "Строка TLS1.3","status == 200, int","PASS"),
        ("TC-PARSE-005","PARSE","ssl_proto = TLSv1.3",
         "Строка TLS1.3","ssl_proto == 'TLSv1.3'","PASS"),
        ("TC-PARSE-006","PARSE","ssl_cipher",
         "Строка TLS1.3","ssl_cipher == 'TLS_AES_256_GCM_SHA384'","PASS"),
        ("TC-PARSE-007","PARSE","request_time",
         "Строка TLS1.3","request_time ≈ 0.015, float","PASS"),
        ("TC-PARSE-008","PARSE","TLS 1.2",
         "Строка TLS1.2","ssl_proto='TLSv1.2', status=302","PASS"),
        ("TC-PARSE-009","PARSE","Редирект 301",
         "Строка redirect","status=301, ssl_proto='-'","PASS"),
        ("TC-PARSE-010","PARSE","Пустая строка → None",
         "''","None","PASS"),
        ("TC-PARSE-011","PARSE","Неверный формат → None",
         "'это не лог'","None","PASS"),
        ("TC-PARSE-012","PARSE","source='file'",
         "Строка + source='file'","source == 'file'","PASS"),
        ("TC-PARSE-013","PARSE","source='syslog'",
         "Строка + source='syslog'","source == 'syslog'","PASS"),
        ("TC-PARSE-014","PARSE","Syslog-префикс",
         "RFC-3164 префикс + строка","Поля распознаны корректно","PASS"),
        ("TC-PARSE-015","PARSE","Все поля присутствуют",
         "Строка TLS1.3","14 обязательных полей в dict","PASS"),
        ("TC-DB-001","DB","Создание таблицы",
         "init_db()","Таблица access_events существует","PASS"),
        ("TC-DB-002","DB","Создание индекса",
         "init_db()","Индекс idx_recorded_at существует","PASS"),
        ("TC-DB-003","DB","save_event записывает",
         "Событие TLS1.3","Строка в таблице, ssl_proto='TLSv1.3'","PASS"),
        ("TC-DB-004","DB","5 событий → +5 строк",
         "5x save_event","count увеличивается на 5","PASS"),
        ("TC-DB-005","DB","Идемпотентность init_db",
         "2x init_db()","Ровно 1 таблица access_events","PASS"),
        ("TC-API-001","API","health 200",
         "GET /api/v1/health","HTTP 200","PASS"),
        ("TC-API-002","API","health status ok",
         "GET /api/v1/health","JSON: status='ok'","PASS"),
        ("TC-API-003","API","logs 200",
         "GET /api/v1/logs","HTTP 200","PASS"),
        ("TC-API-004","API","logs список",
         "GET /api/v1/logs","JSON-массив","PASS"),
        ("TC-API-005","API","Фильтр по TLS",
         "GET /api/v1/logs?protocol=TLSv1.3","Только TLSv1.3","PASS"),
        ("TC-API-006","API","tls stats 200",
         "GET /api/v1/stats/tls","HTTP 200","PASS"),
        ("TC-API-007","API","tls stats структура",
         "GET /api/v1/stats/tls","by_protocol, by_cipher в JSON","PASS"),
        ("TC-API-008","API","rotate 200",
         "POST /api/v1/rotate","HTTP 200","PASS"),
        ("TC-API-009","API","limit",
         "GET /api/v1/logs?limit=2","≤ 2 записи","PASS"),
        ("TC-API-010","API","offset",
         "GET /api/v1/logs?offset=2","offset_data ≤ all_data","PASS"),
        ("TC-SEC-001","SEC","TLS 1.2 и 1.3 в nginx.conf",
         "nginx.conf","TLSv1.2 и TLSv1.3 присутствуют, 1.0/1.1 отсутствуют","PASS"),
        ("TC-SEC-002","SEC","HSTS заголовок",
         "nginx.conf","Strict-Transport-Security, max-age=31536000","PASS"),
        ("TC-SEC-003","SEC","HTTP→HTTPS редирект",
         "nginx.conf","return 301 https://","PASS"),
        ("TC-SEC-004","SEC","server_tokens off",
         "nginx.conf","server_tokens off","PASS"),
        ("TC-SEC-005","SEC","extended_ssl формат",
         "nginx.conf","extended_ssl, $ssl_protocol, $ssl_cipher","PASS"),
        ("TC-SEC-006","SEC","Непривилегированный user",
         "Dockerfile","USER appuser","PASS"),
        ("TC-SEC-007","SEC","HEALTHCHECK",
         "Dockerfile","HEALTHCHECK присутствует","PASS"),
        ("TC-SEC-008","SEC","internal: true",
         "docker-compose.yml","internal: true в сети","PASS"),
        ("TC-SEC-009","SEC","app без внешних портов",
         "docker-compose.yml","ports: [] у сервиса app","PASS"),
        ("TC-SEC-010","SEC","ssl_session_tickets off",
         "nginx.conf","ssl_session_tickets off","PASS"),
        ("TC-INT-001","INT","Файл→БД",
         "Строка TLS1.3 через file watcher","Событие сохранено в БД","PASS"),
    ],
    font_size=9
)
table_caption("Таблица 2.4 — Спецификация тестов (Test Case Specification, IEEE 829)")

h2("2.5  Результаты выполнения тестов")
body("Тесты выполнены 01.05.2026 в среде Replit (Python 3.10.19, pytest 9.0.3).")
body("Итог выполнения:")

code("============================= test session starts ==============================")
code("platform linux -- Python 3.10.19, pytest-9.0.3")
code("collected 41 items")
code("")
code("tests/test_https_container.py::TestParseLineExtendedSSL::...  15 PASSED")
code("tests/test_https_container.py::TestDatabase::...               5 PASSED")
code("tests/test_https_container.py::TestRestAPI::...               10 PASSED")
code("tests/test_https_container.py::TestHTTPSSecurityRequirements::...  10 PASSED")
code("tests/test_https_container.py::TestIntegrationFileTailing::...     1 PASSED")
code("")
code("============================== 41 passed in 0.99s ==============================")

body("Вывод: все 41 тест-кейс пройден успешно. Все цели тестирования T-OBJ-01 — "
     "T-OBJ-05 достигнуты. Требования к КЖД v2.0 не поставлены под угрозу.")
pb()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 3 — ЗАДАНИЕ 6
# ════════════════════════════════════════════════════════════════════════════
h1("3  ЗАДАНИЕ 6. РАЗМЕЩЕНИЕ ОБНОВЛЁННЫХ ПРОДУКТОВ\nВ СРЕДЕ ЗАКАЗЧИКА")
body("Процесс размещения обновлённых программных продуктов в среде заказчика "
     "(Компания 1) регламентируется п. 6.4.7.3.4 ГОСТ Р ИСО/МЭК 12207–2010 "
     "«Перенос» и выполняется в соответствии со стратегией blue-green deployment.")

h2("3.1  Предварительные условия")
bul("Все тесты из раздела 2 пройдены успешно (41/41).")
bul("Компания 1 подписала акт приёмки результатов валидации (Val-01 — Val-04).")
bul("Получен и установлен TLS-сертификат для доменов Компании 1 (Let's Encrypt или корп. УЦ).")
bul("Резервная копия текущей среды создана и верифицирована.")

h2("3.2  Этапы размещения")

h3("3.2.1  Этап 1. Подготовка артефактов")
body("Компания 2 публикует в Container Registry следующие образы:")
bul("access-logger-app:2.0.0  — образ Django-приложения.")
bul("access-logger-kjd:2.0.0  — образ агента КЖД.")
bul("nginx:1.26-alpine         — официальный образ Nginx (pull from Docker Hub).")
body("Каждый образ снабжается SHA-256 дайджестом и подписывается ключом Компании 2 "
     "(Docker Content Trust). Передаётся пакет артефактов развёртывания:")
bul("docker-compose.yml (производственная конфигурация)")
bul("nginx/nginx.conf (конфигурация TLS)")
bul("Инструкция по развёртыванию (настоящий раздел)")
bul("Файл контрольных сумм SHA-256 для всех артефактов")

h3("3.2.2  Этап 2. Развёртывание в тестовом окружении заказчика")
body("Перед переходом в production выполняется развёртывание в staging-среде Компании 1:")

code("# 2.1 Клонирование репозитория / распаковка артефактов")
code("git clone https://repo.company2.example/kjd-v2.git /opt/kjd")
code("cd /opt/kjd")
code("")
code("# 2.2 Генерация/копирование TLS-сертификата")
code("# Производственная среда: использовать сертификат от УЦ!")
code("cp /path/to/company1.crt certs/server.crt")
code("cp /path/to/company1.key certs/server.key")
code("chmod 600 certs/server.key")
code("")
code("# 2.3 Установка секретов в переменные окружения")
code("export DJANGO_SECRET_KEY=$(cat /run/secrets/django_secret_key)")
code("export DJANGO_ALLOWED_HOSTS=www.company1.example,company1.example")
code("")
code("# 2.4 Запуск стека в staging")
code("docker-compose up -d --build")
code("")
code("# 2.5 Проверка health-check всех контейнеров")
code("docker-compose ps")
code("docker inspect kjd_nginx --format '{{.State.Health.Status}}'")

body("После развёртывания в staging выполняются приёмочные тесты Val-01 — Val-04 "
     "(см. раздел 1.6 первого документа).")

h3("3.2.3  Этап 3. Резервное копирование данных текущей среды")

code("# 3.1 Резервная копия тома данных логов")
code("docker run --rm \\")
code("    -v nginx_logs:/data \\")
code("    -v /backup:/backup \\")
code("    alpine tar czf /backup/logs_$(date +%Y%m%d_%H%M%S).tar.gz /data")
code("")
code("# 3.2 Верификация архива")
code("sha256sum /backup/logs_*.tar.gz > /backup/logs_checksums.sha256")
code("")
code("# 3.3 Резервная копия SQLite-базы КЖД")
code("docker run --rm \\")
code("    -v logger_data:/data \\")
code("    -v /backup:/backup \\")
code("    alpine cp /data/events.db /backup/events_$(date +%Y%m%d_%H%M%S).db")

h3("3.2.4  Этап 4. Переключение production (blue-green)")
body("Переключение выполняется методом blue-green deployment для обеспечения "
     "нулевого времени простоя:")

code("# 4.1 Запуск новых ('green') контейнеров рядом со старыми ('blue')")
code("docker-compose -f docker-compose.yml -f docker-compose.green.yml up -d")
code("")
code("# 4.2 Проверка работоспособности 'green' стека")
code("curl -kI https://www.company1.example/  # Ожидается: HTTP/1.1 200 OK")
code("curl -I  http://www.company1.example/   # Ожидается: HTTP/1.1 301 → HTTPS")
code("")
code("# 4.3 Переключение маршрутизатора/балансировщика на 'green'")
code("# (зависит от инфраструктуры Компании 1: Nginx upstream, HAProxy, AWS ALB)")
code("")
code("# 4.4 Мониторинг: проверка отсутствия ошибок в течение 30 минут")
code("docker-compose logs -f nginx | grep -E 'error|crit|emerg'")
code("")
code("# 4.5 Остановка 'blue' контейнеров после подтверждения стабильности")
code("docker-compose -f docker-compose.blue.yml down")

h3("3.2.5  Этап 5. Пост-размещение")
bul("Верификация TLS-конфигурации внешним инструментом (testssl.sh или "
    "https://www.ssllabs.com/ssltest/ — целевая оценка A или A+).")
bul("Проверка наличия HSTS-заголовка в ответах всех сайтов Компании 1.")
bul("Верификация записи событий в КЖД: GET https://logger.company1.example:8080/api/v1/logs.")
bul("Настройка мониторинга: Prometheus + Grafana для визуализации метрик "
    "из /api/v1/stats/tls.")
bul("Подписание акта ввода в эксплуатацию обеими сторонами.")

h2("3.3  Откат (Rollback)")
body("В случае критического сбоя после переключения выполняется откат:")
bul("Переключение балансировщика обратно на 'blue' контейнеры.")
bul("Восстановление данных журналов из резервной копии.")
bul("Фиксация инцидента и передача отчёта в Компанию 2 для анализа.")
pb()

# ════════════════════════════════════════════════════════════════════════════
# РАЗДЕЛ 4 — ЗАДАНИЕ 7
# ════════════════════════════════════════════════════════════════════════════
h1("4  ЗАДАНИЕ 7. ДОВЕДЕНИЕ СВЕДЕНИЙ О МОДИФИКАЦИИ\nДО ЗАИНТЕРЕСОВАННЫХ СТОРОН")
body("Согласно п. 6.4.7.3.2 и п. 6.4.7.3.3 ГОСТ Р ИСО/МЭК 12207–2010, "
     "организация, выполняющая сопровождение, обязана довести сведения "
     "о модификации программных средств до всех затронутых сторон.")

h2("4.1  Идентификация заинтересованных сторон")
mk_table(
    ["Сторона", "Роль", "Затронутость модификацией"],
    [
        ("Компания 1 (заказчик), IT-директор",
         "Принятие решений", "Высокая: изменяются все точки входа"),
        ("Компания 1, администраторы серверов",
         "Эксплуатация инфраструктуры", "Высокая: новые порты, сертификаты, Docker"),
        ("Компания 1, служба безопасности",
         "Контроль соответствия ИБ", "Высокая: смена протокола, TLS-политика"),
        ("Компания 1, конечные пользователи",
         "Использование веб-сервисов", "Средняя: изменение URL с http на https"),
        ("Компания 1, разработчики ПО",
         "Интеграции с API", "Средняя: необходимость обновить клиентский код"),
        ("Компания 2, команда разработки",
         "Поддержка и обновление", "Высокая: новые компоненты (Nginx, КЖД v2)"),
        ("Регулятор / аудитор",
         "Проверка соответствия", "Низкая: переход к HTTPS соответствует требованиям"),
    ]
)
table_caption("Таблица 4.1 — Заинтересованные стороны")

h2("4.2  Порядок уведомления")

h3("4.2.1  Техническое уведомление (Компания 2 → IT Компании 1)")
body("За 10 рабочих дней до развёртывания Компания 2 направляет Компании 1 "
     "технический бюллетень изменений (Change Advisory Board Notice), включающий:")
bul("Описание изменений: переход на HTTPS, новый компонент Nginx, обновление КЖД.")
bul("Список изменённых портов: добавлены 443/TCP и 514/UDP; порт 80 теперь только redirect.")
bul("Перечень новых Docker-образов и их SHA-256 дайджесты.")
bul("Требования к инфраструктуре: наличие TLS-сертификата, открытие порта 443 на firewall.")
bul("Инструкции по откату на случай сбоя.")
bul("Контактные данные службы поддержки Компании 2.")

code("# Шаблон технического бюллетеня (фрагмент)")
code("Тема: [CHANGE-001] Переход на HTTPS — КЖД v2.0")
code("Дата вступления в силу: 15.05.2026")
code("")
code("РЕЗЮМЕ ИЗМЕНЕНИЙ:")
code("  - Веб-сервисы переведены с HTTP/80 на HTTPS/443")
code("  - Добавлен контейнер Nginx 1.26 LTS (TLS-терминация)")
code("  - КЖД обновлён до v2.0 (Alpine 3.21, расширенный формат журналов)")
code("")
code("ДЕЙСТВИЯ АДМИНИСТРАТОРА:")
code("  1. Открыть порт 443/TCP на firewall")
code("  2. Разместить TLS-сертификат в /opt/kjd/certs/")
code("  3. Обновить мониторинг на probe HTTPS вместо HTTP")
code("  4. Обновить ссылки в документации с http:// на https://")

h3("4.2.2  Уведомление конечных пользователей (Компания 1)")
body("За 5 рабочих дней до перехода конечные пользователи уведомляются "
     "через корпоративный портал/рассылку:")

code("Уважаемые коллеги!")
code("")
code("15 мая 2026 г. адреса веб-сервисов компании изменятся:")
code("  БЫЛО:    http://service.company1.example")
code("  СТАЛО:   https://service.company1.example")
code("")
code("При переходе по старым ссылкам браузер автоматически перенаправит")
code("вас на защищённую версию. Дополнительных действий не требуется.")
code("Значок 🔒 в браузере подтвердит защищённое соединение.")
code("")
code("По вопросам: it-support@company1.example")

h3("4.2.3  Уведомление разработчиков (API-клиенты)")
body("Разработчики, использующие API приложения, уведомляются отдельно "
     "с указанием необходимых технических изменений:")

code("CHANGELOG v2.0.0 — Breaking Changes")
code("=========================================")
code("1. Все API-эндпоинты теперь доступны ТОЛЬКО по HTTPS (порт 443).")
code("   Обновите базовый URL клиента:")
code("   БЫЛО:    http://service.company1.example/api/")
code("   СТАЛО:   https://service.company1.example/api/")
code("")
code("2. Новый эндпоинт КЖД REST API:")
code("   GET  https://logger.company1.example:8080/api/v1/logs")
code("   GET  https://logger.company1.example:8080/api/v1/stats/tls")
code("   POST https://logger.company1.example:8080/api/v1/rotate")
code("")
code("3. Самоподписанные сертификаты в STAGING — добавьте в доверенные:")
code("   curl --cacert staging-ca.crt https://staging.company1.example/")

h2("4.3  Документация об изменениях (CHANGELOG)")
body("В соответствии с п. 6.2.2 ГОСТ Р ИСО/МЭК 12207–2010 (управление "
     "конфигурацией) все изменения фиксируются в файле CHANGELOG.md:")

code("# CHANGELOG")
code("")
code("## [2.0.0] — 2026-05-15")
code("### Добавлено")
code("- Nginx 1.26 LTS как обратный прокси с TLS-терминацией")
code("- КЖД v2.0: поддержка формата extended_ssl, SyslogServer (UDP 514)")
code("- REST API: /api/v1/logs, /api/v1/stats/tls, /api/v1/rotate")
code("- Заголовки HSTS, X-Content-Type-Options, X-Frame-Options")
code("- Docker multi-stage build (Alpine 3.21)")
code("")
code("### Изменено")
code("- Протокол HTTP → HTTPS на всех эндпоинтах")
code("- Формат журнала: NCSA Common → extended_ssl (+TLS поля)")
code("- Базовый образ КЖД: Alpine 3.18 → Alpine 3.21")
code("")
code("### Исправлено")
code("- Уязвимость: передача данных в открытом виде (HTTP без шифрования)")
code("")
code("### Удалено")
code("- Прямой доступ к app-контейнеру извне (порт 80/8000 закрыт)")

h2("4.4  Матрица ответственности за уведомление")
mk_table(
    ["Получатель", "Канал уведомления", "Срок", "Ответственный"],
    [
        ("IT-директор Компании 1",   "Официальное письмо",       "За 15 р.д.", "PM Компании 2"),
        ("Администраторы серверов",   "Тех. бюллетень + встреча", "За 10 р.д.", "DevOps Компании 2"),
        ("Служба ИБ",                "Отчёт по безопасности",    "За 10 р.д.", "Security lead К2"),
        ("Конечные пользователи",    "Корп. рассылка / портал",  "За 5 р.д.",  "IT К1 (по запросу К2)"),
        ("Разработчики (API)",        "CHANGELOG + email",        "За 7 р.д.",  "Tech lead Компании 2"),
        ("Аудитор / регулятор",      "Отчёт об изменениях",      "После ввода","PM Компании 2"),
    ]
)
table_caption("Таблица 4.2 — Матрица ответственности за уведомление сторон")
pb()

# ════════════════════════════════════════════════════════════════════════════
# ЗАКЛЮЧЕНИЕ
# ════════════════════════════════════════════════════════════════════════════
h1("ЗАКЛЮЧЕНИЕ")
body("В результате выполнения заданий 4–7 разработаны и задокументированы "
     "следующие материалы в соответствии с требованиями ГОСТ Р ИСО/МЭК 12207–2010, "
     "ГОСТ Р 56920–2016 и IEEE 829-2008:")
bul("Задание 4: разработан и задокументирован модифицированный программный продукт — "
    "контейнер КЖД v2.0 с поддержкой HTTPS. Созданы: Dockerfile (multi-stage), "
    "docker-compose.yml, nginx/nginx.conf, docker/logger/log_agent.py, "
    "certs/generate_self_signed.sh. Приведены алгоритм действий, алгоритм работы "
    "контейнера и описание блок-схемы по ГОСТ 19.701.")
bul("Задание 5: разработан тестовый набор из 41 тест-кейса (pytest). "
    "Составлены: мастер-план тестирования (MTP-KJD-2026-001), спецификация "
    "процедуры тестирования (TPS-KJD-2026-001) и спецификация тестов (TC-PARSE, "
    "TC-DB, TC-API, TC-SEC, TC-INT). Все 41 тест пройдены успешно.")
bul("Задание 6: описан полный процесс размещения обновлённых продуктов в среде "
    "заказчика (5 этапов, blue-green deployment, процедура отката).")
bul("Задание 7: разработаны процедуры доведения сведений о модификации до "
    "всех заинтересованных сторон: технический бюллетень, уведомление пользователей, "
    "CHANGELOG, матрица ответственности.")
pb()

# ════════════════════════════════════════════════════════════════════════════
# СПИСОК ИСТОЧНИКОВ
# ════════════════════════════════════════════════════════════════════════════
h1("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
sources = [
    "ГОСТ Р ИСО/МЭК 12207–2010. Процессы жизненного цикла программных средств. — М.: Стандартинформ, 2011.",
    "ГОСТ Р 56920–2016. Системная и программная инженерия. Тестирование программного обеспечения. Часть 1. — М.: Стандартинформ, 2016.",
    "IEEE Std 829-2008. IEEE Standard for Software and System Test Documentation. — IEEE, 2008.",
    "ГОСТ 19.701–90. Единая система программной документации. Схемы алгоритмов, программ, данных и систем. — М.: Издательство стандартов, 1991.",
    "Nginx Inc. NGINX 1.26 LTS Documentation [Электронный ресурс]. — URL: https://nginx.org/en/docs/ (дата обращения: 01.05.2026).",
    "Mozilla SSL Configuration Generator [Электронный ресурс]. — URL: https://ssl-config.mozilla.org/ (дата обращения: 01.05.2026).",
    "Python Software Foundation. Flask 3.1 Documentation [Электронный ресурс]. — URL: https://flask.palletsprojects.com/ (дата обращения: 01.05.2026).",
    "Docker, Inc. Dockerfile Best Practices [Электронный ресурс]. — URL: https://docs.docker.com/develop/develop-images/dockerfile_best-practices/ (дата обращения: 01.05.2026).",
    "NIST SP 800-52 Rev. 2. Guidelines for TLS Implementations. — NIST, 2019.",
]
for i, s in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}.\t{s}")
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0); pf.left_indent = Cm(1.25)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in p.runs:
        run.font.name = TNR; run.font.size = Pt(14)

out = "Сопровождение_ПС_ГОСТ12207_Задания4-7.docx"
doc.save(out)
print(f"Saved: {out}")
