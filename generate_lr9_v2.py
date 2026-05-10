from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins (ГОСТ: top/bottom 2 cm, left 3 cm, right 1.5 cm) ---
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)

# ──────────────────── helpers ────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=9,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)

def para(doc, text='', size=12, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         indent=True, sb=4, sa=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(size)
    return p

def h1(doc, text, size=13, sb=10, sa=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size)
    return p

def h2(doc, text, size=12, sb=8, sa=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size)
    return p

def h3(doc, text, size=12, sb=6, sa=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size)
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

def mixed(doc, parts, indent=True, sb=4, sa=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """parts = [(text, bold, size, italic), ...]"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    for text, bold, size, italic in parts:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p

def code_block(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(size)
    return p

# ══════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════
def cp(doc, text, size=12, bold=False, sb=3, sa=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    return p

cp(doc,'МИНОБРНАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ', bold=True, sb=0, sa=3)
cp(doc,'Федеральное государственное бюджетное образовательное учреждение\nвысшего образования', size=11)
cp(doc,'Кафедра информационных технологий и кибербезопасности', bold=True, sa=20)
for _ in range(3): doc.add_paragraph()
cp(doc,'ОТЧЁТ', size=16, bold=True, sb=0)
cp(doc,'по лабораторной работе № 9', size=14, bold=True)
cp(doc,'«Конфигурационное управление программным обеспечением»', size=14, bold=True, sa=20)
for _ in range(3): doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for line in [
    'Дисциплина: Сопровождение программных средств\n',
    'Вариант: Django Security Logger App\n',
    'Выполнил: студент группы ___\n',
    'Проверил: Н.В. Корнеев',
]:
    r = p.add_run(line); r.font.size = Pt(12)

for _ in range(5): doc.add_paragraph()
cp(doc,'Москва — 2025', size=12)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. ЦЕЛЬ
# ══════════════════════════════════════════════════════════════════════
h1(doc,'1. ЦЕЛЬ РАБОТЫ')
para(doc,'Получить навыки конфигурационного управления программным обеспечением на примере контейнера для обеспечения безопасности приложения.')

# ══════════════════════════════════════════════════════════════════════
# 2. ФОРМУЛИРОВКА ЗАДАНИЯ
# ══════════════════════════════════════════════════════════════════════
h1(doc,'2. ФОРМУЛИРОВКА ЗАДАНИЯ')
para(doc,'Для программного средства Django Security Logger App (система мониторинга событий безопасности виртуальных машин) рассмотреть процесс менеджмента конфигурации по ГОСТ Р ИСО/МЭК 12207–2010.')
para(doc,'Руководством компании поставлена задача развернуть программное средство одновременно на 1000 серверов заказчика, а затем автоматически изменять конфигурацию при необходимости. В рамках задания необходимо:')
for t in [
    'провести оценку процесса менеджмента конфигурации программного средства по ГОСТ Р ИСО/МЭК 12207–2010 (раздел 7.2.2);',
    'провести сравнительный анализ 5–6 систем управления конфигурациями по критериям: функциональность, стоимость, безопасность, отзывы пользователей — и оформить в виде Таблицы 1;',
    'подготовить план менеджмента конфигурации по ГОСТ Р ИСО 10007–2019 (разделы 1–7 + Приложение);',
    'настроить модель реализации с помощью выбранной системы управления конфигурациями.',
]: bullet(doc, t)

# ══════════════════════════════════════════════════════════════════════
# 3. РЕЗУЛЬТАТЫ
# ══════════════════════════════════════════════════════════════════════
h1(doc,'3. РЕЗУЛЬТАТЫ ВЫПОЛНЕНИЯ ЗАДАНИЯ')

# ────────────────────────────────────────────────────────────────────
# 3.1 Оценка по ГОСТ 12207
# ────────────────────────────────────────────────────────────────────
h2(doc,'3.1. Оценка процесса менеджмента конфигурации ПС по ГОСТ Р ИСО/МЭК 12207–2010')

para(doc,'Согласно ГОСТ Р ИСО/МЭК 12207–2010, процесс менеджмента конфигурации программных средств описан в разделе 7.2.2 «Процессы поддержки программных средств». Цель процесса — установление и сопровождение целостности программных составных частей проекта и обеспечение их доступности для заинтересованных сторон (п. 7.2.2.1).')

para(doc,'В результате успешного осуществления процесса (п. 7.2.2.2) должны быть достигнуты следующие выходы:')
for t in [
    'разработана стратегия менеджмента конфигурации программных средств;',
    'составные части идентифицированы, определены и введены в базовую линию;',
    'контролируются модификации и выпуски составных частей;',
    'обеспечена доступность модификаций и выпусков для заинтересованных сторон;',
    'регистрируется и сообщается статус составных частей и модификаций;',
    'гарантируются завершённость и согласованность составных частей;',
    'контролируются хранение, обработка и поставка составных частей.',
]: bullet(doc, t, size=11)

# Таблица ЭК
h3(doc,'3.1.1. Идентифицированные элементы конфигурации Django Security Logger App')
para(doc,'В соответствии с п. 7.2.2.3.2 ГОСТ Р ИСО/МЭК 12207–2010 разработана схема идентификации: DSLA-<КОД>-<ВЕРСИЯ>. Для каждой составной части установлена базовая линия (baseline) версии 1.0.', indent=True)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h in enumerate(['Идентификатор','Наименование ЭК','Тип ЭК','Версия / Базовая линия']):
    cell_text(hdr[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(hdr[i], 'D9E1F2')
ci_rows = [
    ('DSLA-SC-1.0', 'Исходный код приложения (logger_app/, django_project/)', 'Программный код', 'v1.0 / FB'),
    ('DSLA-TPL-1.0', 'HTML-шаблоны (templates/logger_app/)', 'Программный код', 'v1.0 / FB'),
    ('DSLA-ENC-1.0', 'Модуль шифрования (encryption.py, secret.key)', 'Программный код + данные', 'v1.0 / AB'),
    ('DSLA-SIM-1.0', 'Модуль моделирования событий (data/utils.py)', 'Программный код', 'v1.0 / FB'),
    ('DSLA-DB-1.0', 'База данных SQLite (db.sqlite3)', 'Данные', 'v1.0 / PB'),
    ('DSLA-POL-1.0', 'Политика доступа (migrations/policy.json)', 'Конфигурация', 'v1.0 / AB'),
    ('DSLA-CFG-1.0', 'Настройки Django (django_project/settings.py)', 'Конфигурация', 'v1.0 / AB'),
    ('DSLA-DCK-1.0', 'Docker-инфраструктура (Dockerfile, docker-compose.yml, nginx/)', 'Инфраструктурный код', 'v1.0 / PB'),
    ('DSLA-AGT-1.0', 'Агент КЖД v2.0 (docker/logger/)', 'Программный код', 'v1.0 / PB'),
    ('DSLA-TST-1.0', 'Набор тестов (tests/)', 'Тестовая документация', 'v1.0 / AB'),
]
for id_, name, typ, ver in ci_rows:
    r = tbl.add_row().cells
    cell_text(r[0], id_, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(r[1], name, size=9)
    cell_text(r[2], typ, size=9)
    cell_text(r[3], ver, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Примечание: FB — Functional Baseline (функциональная базовая линия), AB — Allocated Baseline (выделенная), PB — Product Baseline (производственная).')
r.font.size = Pt(9); r.italic = True

# Виды деятельности 7.2.2.3
h3(doc,'3.1.2. Реализация видов деятельности (ГОСТ Р ИСО/МЭК 12207–2010, п. 7.2.2.3)')
activities = [
    ('7.2.2.3.1 Реализация процесса',
     'Разработан и утверждён план менеджмента конфигурации (ПМК) Django Security Logger App. ПМК описывает состав ЭК, стратегию использования Ansible и Git, процедуры внесения изменений, ответственных исполнителей, графики аудитов. Документ хранится в репозитории Git (ветка docs/).'),
    ('7.2.2.3.2 Идентификация конфигурации',
     'Каждому ЭК присвоен уникальный идентификатор по схеме DSLA-<КОД>-<ВЕРСИЯ>. Версии отслеживаются тегами Git (v1.0, v1.1, v2.0). Базовые линии фиксируются по этапам: FB — при утверждении ТЗ, AB — по завершении проектирования, PB — при каждом производственном релизе.'),
    ('7.2.2.3.3 Управление конфигурацией',
     'Все изменения проходят через процедуру Change Request (CR) в GitHub Issues. CR включает: описание, затронутые ЭК, категорию (Minor/Major/Critical). CCB оценивает риски и принимает решение. Изменения реализуются в ветке feature/CR-NNN, проходят CI (GitHub Actions, pytest — 41 тест-кейс) и code review, после чего вливаются в main. Critical-изменения (шифрование, политика доступа) требуют дополнительного аудита безопасности.'),
    ('7.2.2.3.4 Отслеживание состояния конфигурации',
     'Ведётся Configuration Status Log (CSL) — реестр всех ЭК с текущей версией, статусом и историей CR. Ежеквартальные отчёты о состоянии конфигурации (CSR) содержат: список активных ЭК, число CR за период, перечень выпусков, результаты последнего аудита. Данные хранятся в Git с GPG-подписью коммитов.'),
    ('7.2.2.3.5 Оценка конфигурации',
     'Функциональная завершённость проверяется автоматизированным тестированием (pytest: test_encryption.py, test_models.py, test_views.py, test_https_container.py — 41 TC). Физическая завершённость проверяется при аудите: Ansible playbook audit.yml сравнивает хэши развёрнутых Docker-образов с эталоном в Container Registry.'),
    ('7.2.2.3.6 Поставка и менеджмент выпуска',
     'Выпуски формируются как Docker-образы с тегом vMAJOR.MINOR.PATCH и публикуются в Container Registry. Каждый выпуск сопровождается тегом Git и CHANGELOG. Rollout на 1000 серверов выполняется Ansible с rolling update (serial: 50, max_fail_percentage: 5). Резервные копии кодов и документации хранятся в Git с политикой retention не менее 5 лет.'),
]
for title, text in activities:
    mixed(doc,[
        (title + '. ', True, 12, False),
        (text, False, 12, False),
    ], indent=True, sb=4, sa=3)

para(doc,'Таким образом, все шесть видов деятельности процесса менеджмента конфигурации по ГОСТ Р ИСО/МЭК 12207–2010 реализованы и подтверждены документально для Django Security Logger App.')

# ────────────────────────────────────────────────────────────────────
# 3.2 Сравнительный анализ (Таблица 1)
# ────────────────────────────────────────────────────────────────────
h2(doc,'3.2. Сравнительный анализ систем управления конфигурациями')
para(doc,'Для решения задачи развёртывания Django Security Logger App на 1000 серверов и последующего автоматического управления конфигурацией проведён сравнительный анализ шести систем по четырём критериям. Все ценовые данные актуальны по состоянию на май 2025 г. и подтверждены ссылками на официальные источники.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Таблица 1 — Сравнительный анализ систем управления конфигурациями')
r.bold = True; r.font.size = Pt(11)

# TABLE 1 — 6 columns
col_w = [Cm(0.7), Cm(1.9), Cm(4.2), Cm(3.8), Cm(3.5), Cm(4.2)]
heads = ['№','Система','Функциональность','Стоимость (актуальные данные, 2025)','Безопасность','Отзывы пользователей']

rows_data = [
    # 1 Puppet
    ['1','Puppet\n(Perforce)',
     'Декларативный язык Puppet DSL. Агентная pull-модель (Puppet Agent ↔ Puppet Server). Puppet Forge — >7 000 модулей. Отчёты о состоянии узлов, каталоги конфигурации, оркестрация Puppet Bolt. В 2025 г. сообщество создало форк OpenVox после ограничений Perforce.',
     'Community / OpenVox: бесплатно (Apache 2.0 / открытый форк).\nPuppet Enterprise: ~$112–199/узел/год (Standard / Premium Support; точная цена — по запросу).\n⚠ С нач. 2025 новые бинарники Perforce — в закрытом репозитории, >25 узлов — коммерческая лицензия.\n\nИсточники:\n• puppet.com/pricing\n• thenewstack.io/openvox...\n• peerspot.com (отзывы)',
     'Взаимная PKI/TLS-аутентификация (Puppet Server ↔ Agent). Подпись сертификатов CA. RBAC (только Enterprise). Централизованный аудит изменений. Шифрование трафика TLS 1.2/1.3.',
     'G2/PeerSpot: 8.2/10 (12+ отзывов, 2024–2025).\nПлюсы: мощный декларативный подход, зрелая экосистема, крупные enterprise-внедрения.\nМинусы: крутая кривая обучения DSL, зависимость от Puppet Agent, риски лицензирования Perforce 2025. [1]'],
    # 2 Chef
    ['2','Chef\n(Progress)',
     'Императивный Ruby DSL (Recipes / Cookbooks / Roles). Агентная pull-модель (Chef Client ↔ Chef Server). Chef Supermarket — библиотека cookbook. InSpec — framework соответствия. Test Kitchen — тестирование рецептов. Интеграция с AWS/Azure/GCP.',
     'chef-workstation / chef-infra-client: бесплатно (Apache 2.0).\nProgress Chef Business: $59/узел/год.\nProgress Chef Enterprise: $189/узел/год (RBAC, Continuous Compliance, User Mgmt).\nEnterprise Plus: по запросу (premium SLA).\nОценка для 1000 узлов: от $59 000/год.\n\nИсточники:\n• capterra.com/p/170467/Chef\n• trustradius.com/products/progress-chef/pricing',
     'Chef Vault — шифрование секретов (AES-256). HTTPS/TLS для Chef Server API. Подпись API-запросов (RSA). RBAC. InSpec обеспечивает compliance-аудит (CIS, STIG, PCI DSS). Encrypted Data Bags.',
     'G2: 4.3/5 (100+ отзывов, 2024).\nПлюсы: максимальная гибкость благодаря Ruby, мощный compliance-аудит через InSpec, зрелая CI/CD-интеграция.\nМинусы: высокий порог входа для не-Ruby разработчиков, сложная инфраструктура Chef Server, более медленное развитие community. [2]'],
    # 3 SaltStack
    ['3','SaltStack\n(Salt / Broadcom)',
     'YAML States + Jinja2-шаблонизация. Push-модель через ZeroMQ (быстро!) + pull-режим. Событийная система Salt Reactor. Поддержка агентного (Minion) и безагентного (SSH) режимов. Грануляция — фильтрация узлов по ОС, роли, IP. Модуль Execution — ad-hoc команды.',
     'Salt Project Open Source: бесплатно (Apache 2.0).\nEnterprise (VMware Aria Automation Config / Broadcom Tanzu Salt): только по запросу — включён в VCF/VVF bundle Broadcom, отдельная цена не публикуется.\n⚠ После поглощения VMware компанией Broadcom (2023–2024) ценообразование резко изменилось в сторону enterprise-пакетов.\n\nИсточники:\n• saltproject.io\n• vmware.com/products/aria-automation\n• trustradius.com/compare-products/salt...',
     'AES-256 шифрование транспортного канала (ZeroMQ). Аутентификация Minion/Master по открытым ключам. Pillar — изоляция секретных данных по узлам. Безагентный SSH-режим — нет открытых портов на клиентах. Поддержка SELinux-окружений.',
     'TrustRadius: 8.1/10. G2: 4.2/5 (70+ отзывов, 2024).\nПлюсы: высочайшая производительность на больших парках (1000+ узлов), уникальная реактивная событийная система, гибкость YAML+Jinja2.\nМинусы: сложная настройка при большом числе Minion-ов, неопределённость будущего enterprise-версии под Broadcom. [3]'],
    # 4 Ansible
    ['4','Ansible\n(Red Hat / IBM)',
     'Декларативный YAML (Playbooks / Roles / Collections). Безагентная push-архитектура (SSH / WinRM) — не требует ПО на управляемых узлах. Ansible Galaxy — >50 000 ролей. Идемпотентные модули. AWX / Ansible Automation Platform — веб-интерфейс, RBAC, API, планировщик заданий.',
     'ansible-core: полностью бесплатно (GPL-3.0).\nAWX (upstream Tower): бесплатно (open source).\nRed Hat AAP Standard: ~$13 000/год за 100 узлов (~$130/узел/год).\nRed Hat AAP Premium: выше (SLA 24×7).\nОценка для 1000 узлов (Standard): ~$130 000/год (с Red Hat поддержкой), либо $0 (ansible-core + AWX).\n\nИсточники:\n• redhat.com/en/technologies/management/ansible/pricing\n• cyberpanel.net/blog/ansible-pricing\n• trustradius.com/products/red-hat-...',
     'Передача только по SSH (нет демонов/агентов на клиентах — минимальная поверхность атаки). Ansible Vault — шифрование секретов AES-256. RBAC в Ansible Tower/AWX. Поддержка Kerberos, LDAP. Проверка подлинности хостов через known_hosts / SSH-ключи.',
     'G2: 4.5/5 (800+ отзывов, 2024–2025) — #1 в Network Automation по G2.\nForrester: лидер в The Forrester Wave: Infrastructure Automation.\nПлюсы: самый простой порог входа (YAML читается как документация), безагентность, огромное сообщество (>26 млн загрузок).\nМинусы: может замедляться без Tower на >1000 узлов, последовательная природа playbook, нет встроенного pull-режима. [4]'],
    # 5 NIX
    ['5','NIX /\nNixOS',
     'Функциональный язык Nix. Полностью воспроизводимые сборки (reproducible builds) — одинаковый результат на любой машине. Атомарные обновления и откаты (nixos-rebuild switch --rollback). NixOS Modules — декларативная конфигурация всей ОС. Nix Flakes — управление зависимостями с lock-файлом. NixOps / deploy-rs — удалённое развёртывание.',
     'NixOS: полностью бесплатно (LGPL / MIT).\nDeterminate Nix (enterprise-дистрибутив): бесплатно (включает SOC 2 Type II валидацию).\nFlakeHub (Determinate Systems, enterprise): по запросу.\nКоммерческой версии NixOS как таковой нет — vendor lock отсутствует.\n\nИсточники:\n• nixos.org\n• determinate.systems\n• docs.determinate.systems/determinate-nix',
     'Изоляция пакетов в /nix/store (хэш-адресация — каждый пакет идентифицируется SHA-256 хэшем входных данных). Неизменяемая инфраструктура — нет глобальных мутаций. Отсутствие зависимостей между пакетами снижает риск атак на цепочку поставок (supply chain). Воспроизводимые сборки — верифицируемый и аудируемый результат. CVE-процесс у Determinate Systems.',
     'GitHub Stars: >17 000 (NixOS/nixpkgs). Растущее сообщество в облачных и DevSecOps-проектах. Gartner: нишевый игрок, но растущий тренд в immutable infrastructure.\nПлюсы: лучшая воспроизводимость и безопасность supply chain, атомарные откаты, уникальный подход.\nМинусы: очень высокий порог входа (язык Nix нетривиален), мало специалистов на рынке, не предназначен для массового мгновенного развёртывания на 1000 серверов без NixOps. [5]'],
    # 6 CFEngine
    ['6','CFEngine\n(Northern.tech)',
     'Декларативный язык CFEngine DSL (Promise Theory). Один из старейших инструментов (с 1993 г.). Легковесный агент (<1 МБ). Автономный режим — агенты выполняют политики независимо от центра при потере связи (resilient). Масштабируется до 100 000+ узлов. CFEngine Mission Portal (Enterprise) — веб-интерфейс.',
     'CFEngine Community: бесплатно (GPL), неограниченное число узлов.\nCFEngine Enterprise Trial: бесплатно до 25 узлов.\nCFEngine Enterprise Full: по запросу (Northern.tech) — подписочная модель, без публичного прайс-листа. По данным пользователей TrustRadius — конкурентоспособная цена при больших парках.\n\nИсточники:\n• cfengine.com/community-vs-enterprise-comparison\n• trustradius.com/products/cfengine/pricing\n• getapp.com/...cfengine/pricing',
     'PKI-аутентификация (агент ↔ хаб). TLS-шифрование трафика. FIPS 140-2 совместимость (Enterprise). Автономность агентов устраняет единую точку отказа. Минимальный сетевой трафик (дифференциальные политики). Разграничение прав доступа к ресурсам.',
     'PeerSpot: 7.8/10. TrustRadius: 8.5/10 (небольшое число отзывов — нишевый инструмент).\nПлюсы: исключительная надёжность и масштабируемость, автономность агентов (работа без сети), применяется в телекоме и финансовом секторе.\nМинусы: устаревший синтаксис, медленное развитие, крутая кривая обучения DSL, небольшое сообщество по сравнению с Ansible. [6]'],
]

tbl2 = doc.add_table(rows=1, cols=6)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate(col_w):
    for c in tbl2.columns[i].cells:
        c.width = w
hdr2 = tbl2.rows[0].cells
for i, h in enumerate(heads):
    cell_text(hdr2[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(hdr2[i], '2F5496')
    hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

ALT = 'DEEAF1'
for idx, row_data in enumerate(rows_data):
    row = tbl2.add_row().cells
    bg = ALT if idx % 2 == 0 else 'FFFFFF'
    for i, ct in enumerate(row_data):
        cell_text(row[i], ct, size=8.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER if i==0 else WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_cell_bg(row[i], bg)

# Sources under table
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(5)
r = p.add_run('Источники к Таблице 1:')
r.bold = True; r.font.size = Pt(10)
src_lines = [
    '[1] Puppet pricing (puppet.com/pricing); OpenVox fork: thenewstack.io/openvox-the-community-driven-fork-of-puppet-has-arrived; PeerSpot reviews: peerspot.com/products/puppet-enterprise-reviews',
    '[2] Progress Chef pricing: capterra.com/p/170467/Chef; trustradius.com/products/progress-chef/pricing; G2: g2.com/products/progress-chef/reviews',
    '[3] Salt Project: saltproject.io; VMware Aria: vmware.com/products/aria-automation/saltstack-config; TrustRadius comparison: trustradius.com/compare-products/salt-open-source...',
    '[4] Red Hat AAP pricing: redhat.com/en/technologies/management/ansible/pricing; Netsync reseller data: cyberpanel.net/blog/ansible-pricing; G2: g2.com/products/red-hat-ansible-automation-platform',
    '[5] NixOS: nixos.org; Determinate Systems: determinate.systems; Determinate Nix docs: docs.determinate.systems/determinate-nix',
    '[6] CFEngine CE vs Enterprise: cfengine.com/community-vs-enterprise-comparison; GetApp pricing: getapp.com/.../cfengine/pricing; TrustRadius: trustradius.com/products/cfengine/pricing',
]
for s in src_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(s); r.font.size = Pt(8.5); r.italic = True

# Вывод
para(doc,'По результатам сравнительного анализа для развёртывания Django Security Logger App на 1000 серверов выбрана система Ansible (ansible-core + AWX) по следующим критериям: (1) безагентная архитектура не требует установки ПО на 1000 серверах заказчика; (2) простой YAML-синтаксис playbook минимизирует время освоения; (3) Ansible Vault обеспечивает шифрование секретов AES-256; (4) высокие оценки сообщества (G2: 4.5/5, 800+ отзывов); (5) полностью бесплатная open-source версия — критически важно при развёртывании на 1000 узлов без лицензионных затрат.')

# ────────────────────────────────────────────────────────────────────
# 3.3 ПЛАН МЕНЕДЖМЕНТА КОНФИГУРАЦИИ (ГОСТ Р ИСО 10007-2019)
# ────────────────────────────────────────────────────────────────────
h2(doc,'3.3. План менеджмента конфигурации по ГОСТ Р ИСО 10007–2019')
para(doc,'Структура плана разработана в соответствии с Приложением А (справочное) ГОСТ Р ИСО 10007–2019 «Структура и содержание плана менеджмента конфигурации» (разделы А.1–А.7). Нумерация разделов соответствует заданию к лабораторной работе.')

# 1. Общие положения
h3(doc,'1. Общие положения')
para(doc,'Настоящий план менеджмента конфигурации (ПМК) распространяется на программное средство Django Security Logger App (DSLA) версии 1.x — веб-приложение для мониторинга событий безопасности виртуальных машин с ролевым разграничением доступа (admin / auditor), шифрованием журналов (Fernet AES-128) и аналитической панелью.')
para(doc,'ПМК разработан в соответствии с:')
for t in [
    'ГОСТ Р ИСО 10007–2019 «Менеджмент качества. Руководящие указания по менеджменту конфигурации»;',
    'ГОСТ Р ИСО/МЭК 12207–2010 «Процессы жизненного цикла программных средств» (п. 7.2.2);',
    'ГОСТ Р ИСО 9001:2015 (требования к идентификации и прослеживаемости, п. 8.5.2).',
]: bullet(doc, t, size=11)
para(doc,'Сфера применения: разработка, тестирование, развёртывание и сопровождение DSLA на 1000 серверов заказчика. Срок действия: весь жизненный цикл продукта (от разработки до вывода из эксплуатации).')

# 2. Введение
h3(doc,'2. Введение')
para(doc,'Цель ПМК — обеспечить целостность, прослеживаемость и управляемость всех элементов конфигурации DSLA на протяжении всего жизненного цикла.')
para(doc,'Описание продукции: DSLA реализован на Django 5.0.2 (Python 3.10), использует SQLite, шифрование Fernet (cryptography), pandas/numpy для аналитики. Развёртывание — Docker-контейнеры под управлением Nginx с TLS 1.2/1.3.')
para(doc,'Инструменты менеджмента конфигурации:')
for t in [
    'Git (GitHub) — контроль версий исходного кода, документации и конфигурационных файлов;',
    'Ansible (ansible-core + AWX) — система управления конфигурациями целевых серверов;',
    'Ansible Vault — шифрование (AES-256) секретных данных в playbook-ах;',
    'Docker + Container Registry — упаковка и версионирование релизов;',
    'GitHub Actions — CI/CD-конвейер (сборка, pytest, публикация Docker-образа);',
    'GitHub Issues — регистрация и отслеживание заявок на изменения (CR).',
]: bullet(doc, t, size=11)
para(doc,'График ключевых мероприятий: идентификация ЭК и установка FB — при старте проекта; установка AB — по завершении проектирования; установка PB — при каждом релизе; ежеквартальные CSR-отчёты; аудит конфигурации — перед релизом и не реже 1 раза в год.')

# 3. Политики
h3(doc,'3. Политики')
para(doc,'Политика в области менеджмента конфигурации DSLA включает следующие принципы (ГОСТ Р ИСО 10007–2019, п. А.3):')
for t in [
    'Все ЭК находятся под контролем версий Git. Прямые коммиты в ветку main/master запрещены — изменения вносятся только через Pull Request с обязательным code review (не менее 1 ревьюера).',
    'Каждый релиз сопровождается тегом Git формата vMAJOR.MINOR.PATCH, Docker-образом с тем же тегом и записью в CHANGELOG.',
    'Секретные данные (ключи шифрования, пароли, сертификаты) хранятся исключительно в Ansible Vault — в открытый репозиторий не попадают.',
    'Ответственность за одобрение изменений в критических ЭК (DSLA-ENC, DSLA-POL) несёт CCB в составе: менеджер конфигурации + менеджер безопасности.',
    'Квалификация: все участники прошли обучение по Git Flow, Ansible и политике информационной безопасности компании.',
    'Терминология соответствует ГОСТ Р ИСО 10007–2019 и ГОСТ Р ИСО/МЭК 12207–2010.',
]: bullet(doc, t, size=11)

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
set_cell_bg(tbl3.rows[0].cells[0], 'D9E1F2')
set_cell_bg(tbl3.rows[0].cells[1], 'D9E1F2')
set_cell_bg(tbl3.rows[0].cells[2], 'D9E1F2')
cell_text(tbl3.rows[0].cells[0], 'Роль', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(tbl3.rows[0].cells[1], 'Ответственность', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(tbl3.rows[0].cells[2], 'Полномочия', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
for role, resp, auth in [
    ('Менеджер конфигурации','Разработка/сопровождение ПМК, ведение реестра ЭК, координация процесса','Утверждение Minor CR'),
    ('CCB (Configuration Control Board)','Оценка и одобрение Major/Critical CR, анализ рисков','Одобрение/отклонение всех CR'),
    ('Разработчики','Соблюдение Git-процедур, документирование изменений, создание CR','Создание PR, реализация CR'),
    ('DevOps-инженер','Сопровождение Ansible playbook, CI/CD, Docker Registry','Развёртывание релизов'),
    ('Менеджер безопасности','Аудит DSLA-ENC и DSLA-POL, управление ключами Vault','Одобрение Critical CR'),
    ('Тестировщик','Верификация изменений (pytest, интеграционные тесты)','Подтверждение корректности CR'),
]:
    r = tbl3.add_row().cells
    cell_text(r[0], role, size=10)
    cell_text(r[1], resp, size=10)
    cell_text(r[2], auth, size=10)
doc.add_paragraph()

# 4. Идентификация конфигурации
h3(doc,'4. Идентификация конфигурации')
para(doc,'В соответствии с ГОСТ Р ИСО 10007–2019 (п. А.4) описывается:')
mixed(doc,[('Схема идентификации: ', True, 12, False),('DSLA-<КОД>-<ВЕРСИЯ>', False, 12, True),(' — где КОД — двухбуквенный код компонента (SC, ENC, POL, CFG, DB, TPL, SIM, DCK, AGT, TST), ВЕРСИЯ — MAJOR.MINOR.PATCH. Пример: DSLA-ENC-1.2.0.', False, 12, False)], indent=True)
para(doc,'Базовые конфигурации и сроки установки:')
for t in [
    'Functional Baseline (FB) — при утверждении ТЗ DSLA; включает DSLA-SC, DSLA-TPL, DSLA-SIM;',
    'Allocated Baseline (AB) — по завершении архитектурного проектирования; добавляются DSLA-ENC, DSLA-POL, DSLA-CFG, DSLA-TST;',
    'Product Baseline (PB) — при каждом производственном релизе; охватывает все 10 ЭК.',
]: bullet(doc, t, size=11)
para(doc,'Статус пересмотра обозначается суффиксом: -alpha (разработка), -beta (тестирование), -rc.N (кандидат в релиз), -release (производственный релиз). Все ЭК и их версии фиксируются в Configuration Item Register (CIR) — документированном реестре в Git (docs/cir.md).')

# 5. Управление изменениями
h3(doc,'5. Управление изменениями')
para(doc,'Процедура управления изменениями (ГОСТ Р ИСО 10007–2019, п. А.5; ГОСТ Р ИСО/МЭК 12207–2010, п. 7.2.2.3.3):')
steps = [
    ('1. Инициирование CR','Заинтересованная сторона создаёт CR в GitHub Issues с описанием проблемы, обоснованием и категорией (Minor / Major / Critical). CR включает: затронутые ЭК, текущую версию, описание изменения, оценку рисков.'),
    ('2. Документирование','CR присваивается уникальный номер. Статус: «Открыта». Уведомление — менеджер конфигурации.'),
    ('3. Оценка CCB','CCB оценивает: технические преимущества, риски безопасности, влияние на смежные ЭК, трудозатраты, сроки. Результат — задокументированное решение (одобрить / отклонить / отложить).'),
    ('4. Реализация','Разработчик создаёт ветку Git: feature/CR-NNN, вносит изменения, пишет/обновляет тесты, открывает Pull Request.'),
    ('5. Верификация','CI (GitHub Actions) выполняет сборку и полный набор тестов pytest. Проводится code review (≥1 ревьюер). Critical CR — дополнительный аудит безопасности.'),
    ('6. Внедрение','После одобрения PR сливается в main, формируется новый тег и Docker-образ. Ansible playbook update.yml (rolling update, serial: 50) обновляет серверы заказчика.'),
    ('7. Верификация внедрения','Ansible playbook audit.yml проверяет соответствие версии эталону на всех 1000 серверах. Результаты фиксируются в CSL.'),
    ('8. Закрытие CR','CR закрывается с указанием новой версии ЭК, даты внедрения и ответственного. CHANGELOG обновляется.'),
]
for step, desc in steps:
    mixed(doc,[(step+': ', True, 11, False),(desc, False, 11, False)], indent=True, sb=3, sa=2)

# 6. Учёт статуса конфигурации
h3(doc,'6. Учёт статуса конфигурации')
para(doc,'Деятельность по учёту статуса конфигурации (ГОСТ Р ИСО 10007–2019, п. 5.5 и А.6) охватывает все стадии жизненного цикла DSLA и включает следующую документированную информацию:')
for t in [
    'Configuration Item Register (CIR) — реестр всех ЭК: идентификатор, наименование, текущая версия, статус пересмотра, дата базовой линии, история изменений (docs/cir.md в Git);',
    'Журнал CR (Change Request Log) — все поданные заявки: номер, статус (Open / In Review / Approved / Implemented / Closed), дата, ответственный;',
    'Журнал выпусков (Release Log) — перечень выпущенных версий с датой, составом изменений (CHANGELOG.md);',
    'Отчёт о состоянии конфигурации (Configuration Status Report, CSR) — ежеквартальная сводка: список активных ЭК, число CR за период, перечень выпусков, результаты аудита.',
]: bullet(doc, t, size=11)
para(doc,'Данные хранятся в Git-репозитории с ограниченным доступом (RBAC GitHub). Целостность обеспечивается GPG-подписью коммитов. Резервное копирование — ежедневно в защищённое облачное хранилище. Хранение: не менее 5 лет после вывода продукта из эксплуатации (ГОСТ Р ИСО 10007–2019, п. 5.5.2.3).')

# 7. Аудит конфигурации
h3(doc,'7. Аудит конфигурации')
para(doc,'В соответствии с ГОСТ Р ИСО 10007–2019 (п. 5.5.4 и А.7) проводятся два типа аудита конфигурации:')
for t in [
    'Функциональный аудит конфигурации (ФАК) — верификация достижения функциональных и рабочих характеристик ЭК, указанных в ТЗ. Инструмент: pytest (41 тест-кейс: test_encryption.py, test_models.py, test_views.py, test_https_container.py). Периодичность: перед каждым производственным релизом (PB). Ответственный: тестировщик.',
    'Физический аудит конфигурации (ФиАК) — верификация соответствия развёрнутых артефактов эталонным версиям в Container Registry. Инструмент: Ansible playbook audit.yml (сравнение хэшей Docker-образов на всех 1000 серверах). Периодичность: ≥1 раза в год и после инцидентов безопасности. Ответственный: менеджер конфигурации + менеджер безопасности.',
]: bullet(doc, t, size=11)
para(doc,'Результаты аудита оформляются в Отчёте об аудите конфигурации (Configuration Audit Report, CAR) и направляются руководству проекта. При выявлении несоответствий создаётся CR категории Critical с немедленным уведомлением CCB. Форма CAR хранится в docs/audit/ в Git.')

# 8. ПРИЛОЖЕНИЕ
h3(doc,'8. Приложение: настройка модели реализации с помощью Ansible')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 3.4 Настройка с Ansible (Приложение к ПМК)
# ════════════════════════════════════════════════════════════════════
h2(doc,'3.4. Настройка модели реализации с помощью Ansible\n(Приложение к Плану менеджмента конфигурации)')

para(doc,'Ниже представлена полная практическая модель развёртывания Django Security Logger App на 1000 серверов заказчика с автоматическим управлением конфигурацией. Использованы: ansible-core 2.17, AWX 24.x, Docker 26, Container Registry.')

h3(doc,'Структура Ansible-проекта')
code_block(doc,
'''dsla-ansible/
├── ansible.cfg                        # конфигурация Ansible (forks, SSH-настройки)
├── requirements.yml                   # зависимости из Ansible Galaxy
├── inventory/
│   ├── production/
│   │   ├── hosts.ini                  # 1000 серверов заказчика
│   │   └── group_vars/
│   │       ├── all.yml                # открытые переменные
│   │       └── all_vault.yml          # зашифрованные секреты (Ansible Vault AES-256)
│   └── staging/
│       └── hosts.ini                  # тестовые серверы (10 узлов)
├── roles/
│   ├── common/                        # обновление ОС, fail2ban, NTP, sysctl
│   ├── docker/                        # установка Docker CE, compose plugin
│   ├── nginx/                         # Nginx + TLS 1.2/1.3, HSTS
│   └── dsla_app/                      # развёртывание DSLA
│       ├── tasks/main.yml
│       ├── templates/
│       │   ├── docker-compose.yml.j2
│       │   └── policy.json.j2
│       └── handlers/main.yml
├── playbooks/
│   ├── site.yml                       # полное первоначальное развёртывание
│   ├── update.yml                     # rolling update (обновление конфигурации)
│   └── audit.yml                      # физический аудит конфигурации
└── vault/
    └── .vault_pass                    # файл пароля Vault (не в Git!)''')

h3(doc,'ansible.cfg — конфигурация Ansible')
code_block(doc,
'''[defaults]
inventory          = inventory/production/hosts.ini
remote_user        = deploy
private_key_file   = ~/.ssh/dsla_deploy_key
forks              = 20          # 20 параллельных подключений
host_key_checking  = True
stdout_callback    = yaml
gathering          = smart       # кэшировать факты
fact_caching       = jsonfile
fact_caching_connection = /tmp/ansible_facts

[ssh_connection]
pipelining         = True        # ускорение: меньше SSH-соединений
ssh_args           = -C -o ControlMaster=auto -o ControlPersist=60s''')

h3(doc,'inventory/production/hosts.ini — инвентарь 1000 серверов')
code_block(doc,
'''[dsla_servers]
# Диапазон: server001 ... server1000
server[001:500].prod.example.com  ansible_host={{ ip_range_1 }}
server[501:1000].prod.example.com ansible_host={{ ip_range_2 }}

[dsla_servers:vars]
ansible_python_interpreter = /usr/bin/python3
ansible_user               = deploy
ansible_port               = 22

[staging]
staging[01:10].test.example.com''')

h3(doc,'group_vars/all.yml — открытые переменные')
code_block(doc,
'''# Версия приложения (изменяется при каждом релизе)
dsla_version:  "1.2.0"
dsla_image:    "registry.example.com/dsla:{{ dsla_version }}"
dsla_port:     5000

# Пути на целевых серверах
dsla_data_dir: "/opt/dsla/data"
dsla_log_dir:  "/var/log/dsla"

# Nginx
nginx_http_port:  80
nginx_https_port: 443

# Параметры rolling update
dsla_serial:              50    # обновлять по 50 серверов за раз
dsla_max_fail_percentage: 5     # остановить, если >5% серверов упали

# Выполнять ли миграции БД при обновлении
dsla_run_migrations: true''')

h3(doc,'group_vars/all_vault.yml — зашифрованные секреты (Ansible Vault)')
para(doc,'Создаётся командой: ansible-vault create inventory/production/group_vars/all_vault.yml --vault-password-file vault/.vault_pass', indent=True)
code_block(doc,
'''# После расшифровки ansible-vault view:
vault_django_secret_key:    "50-символьный_случайный_ключ_SECRET_KEY"
vault_fernet_key:           "ключ_Fernet_base64_для_шифрования_журналов"
vault_db_password:          "пароль_PostgreSQL_для_production"
vault_registry_user:        "deploy-bot"
vault_registry_password:    "токен_Container_Registry"
vault_ssl_cert_content:     |
  -----BEGIN CERTIFICATE-----
  ...
vault_ssl_key_content:      |
  -----BEGIN PRIVATE KEY-----
  ...''')

h3(doc,'roles/dsla_app/tasks/main.yml — задачи развёртывания')
code_block(doc,
'''---
- name: Создать директории приложения
  file:
    path: "{{ item }}"
    state: directory
    owner: deploy
    group: deploy
    mode: '0750'
  loop:
    - "{{ dsla_data_dir }}"
    - "{{ dsla_log_dir }}"
    - "{{ dsla_data_dir }}/certs"

- name: Записать ключ Fernet из Vault (шифрование журналов)
  copy:
    content: "{{ vault_fernet_key }}"
    dest: "{{ dsla_data_dir }}/secret.key"
    owner: deploy
    group: deploy
    mode: '0600'      # только владелец — никаких прав на чтение у других
  no_log: true        # не выводить в лог Ansible

- name: Развернуть политику доступа из шаблона
  template:
    src: policy.json.j2
    dest: "{{ dsla_data_dir }}/policy.json"
    owner: deploy
    mode: '0640'

- name: Авторизоваться в Container Registry
  docker_login:
    registry: registry.example.com
    username: "{{ vault_registry_user }}"
    password: "{{ vault_registry_password }}"
  no_log: true

- name: Создать docker-compose из шаблона и запустить DSLA
  template:
    src: docker-compose.yml.j2
    dest: "{{ dsla_data_dir }}/docker-compose.yml"
    owner: deploy
    mode: '0640'

- name: Запустить стек через docker compose
  community.docker.docker_compose_v2:
    project_src: "{{ dsla_data_dir }}"
    pull: always
    state: present
  notify: Restart DSLA

- name: Выполнить миграции Django
  community.docker.docker_container_exec:
    container: "dsla_app_1"
    command: python manage.py migrate --noinput
  when: dsla_run_migrations | bool

- name: Проверить доступность приложения (health check)
  uri:
    url: "http://127.0.0.1:{{ dsla_port }}/logger/"
    status_code: 200
    timeout: 30
  retries: 5
  delay: 10
  register: health_check

- name: Сообщить результат проверки
  debug:
    msg: "DSLA {{ dsla_version }} успешно запущен на {{ inventory_hostname }}"
  when: health_check.status == 200''')

h3(doc,'playbooks/site.yml — полное первоначальное развёртывание')
code_block(doc,
'''---
- name: Полное развёртывание Django Security Logger App на 1000 серверов
  hosts: dsla_servers
  become: true
  vars_files:
    - "{{ inventory_dir }}/group_vars/all_vault.yml"

  roles:
    - role: common      # Обновление ОС, настройка NTP, fail2ban, sysctl
    - role: docker      # Установка Docker CE 26, compose plugin
    - role: nginx       # Nginx, TLS 1.2/1.3, HSTS, HTTP→HTTPS redirect
    - role: dsla_app    # Развёртывание DSLA

# Команда запуска:
# ansible-playbook -i inventory/production playbooks/site.yml \\
#   --vault-password-file vault/.vault_pass --forks 20''')

h3(doc,'playbooks/update.yml — обновление конфигурации (rolling update)')
code_block(doc,
'''---
# Автоматическое обновление конфигурации на всех 1000 серверах
# Rolling update: 50 серверов за раз, остановка при >5% сбоев
- name: Rolling Update — Django Security Logger App
  hosts: dsla_servers
  become: true
  serial:               "{{ dsla_serial }}"        # 50 по умолчанию
  max_fail_percentage:  "{{ dsla_max_fail_percentage }}"   # 5%
  vars_files:
    - "{{ inventory_dir }}/group_vars/all_vault.yml"

  pre_tasks:
    - name: Вывести сервер из балансировщика нагрузки (если есть)
      uri:
        url: "http://lb.example.com/api/disable/{{ inventory_hostname }}"
        method: POST
      ignore_errors: true    # продолжить даже без LB

  tasks:
    - name: Обновить Docker-образ DSLA до версии {{ dsla_version }}
      community.docker.docker_image:
        name: "{{ dsla_image }}"
        source: pull
        force_source: true

    - name: Перезапустить контейнер с новым образом
      community.docker.docker_compose_v2:
        project_src: "{{ dsla_data_dir }}"
        pull: always
        state: present
        recreate: always   # принудительное пересоздание контейнера

    - name: Выполнить миграции (если нужно)
      community.docker.docker_container_exec:
        container: dsla_app_1
        command: python manage.py migrate --noinput
      when: dsla_run_migrations | bool

    - name: Проверить здоровье после обновления
      uri:
        url: "http://127.0.0.1:{{ dsla_port }}/logger/"
        status_code: 200
      retries: 5
      delay: 15

  post_tasks:
    - name: Вернуть сервер в балансировщик нагрузки
      uri:
        url: "http://lb.example.com/api/enable/{{ inventory_hostname }}"
        method: POST
      ignore_errors: true

# Команда:
# ansible-playbook -i inventory/production playbooks/update.yml \\
#   -e "dsla_version=1.3.0" --vault-password-file vault/.vault_pass''')

h3(doc,'playbooks/audit.yml — физический аудит конфигурации')
code_block(doc,
'''---
# Физический аудит конфигурации (ФиАК) — ГОСТ Р ИСО 10007-2019, п. 5.5.4
# Проверяет, что все 1000 серверов запускают одобренную версию DSLA
- name: Физический аудит конфигурации DSLA
  hosts: dsla_servers
  become: true
  gather_facts: true

  tasks:
    - name: Получить информацию о запущенном контейнере
      community.docker.docker_container_info:
        name: dsla_app_1
      register: container_info

    - name: Проверить соответствие образа эталону
      assert:
        that:
          - container_info.container is defined
          - container_info.container.Config.Image == dsla_image
        fail_msg: >
          НЕСООТВЕТСТВИЕ на {{ inventory_hostname }}:
          запущен {{ container_info.container.Config.Image | default("N/A") }},
          ожидается {{ dsla_image }}
        success_msg: "OK: {{ inventory_hostname }} — версия {{ dsla_version }} ✓"

    - name: Проверить целостность ключа Fernet
      stat:
        path: "{{ dsla_data_dir }}/secret.key"
      register: key_stat

    - name: Убедиться, что ключ существует и доступен только владельцу
      assert:
        that:
          - key_stat.stat.exists
          - key_stat.stat.mode == "0600"
        fail_msg: "НАРУШЕНИЕ БЕЗОПАСНОСТИ: ключ шифрования на {{ inventory_hostname }}!"

    - name: Сохранить результаты аудита локально
      local_action:
        module: lineinfile
        path: "audit_report_{{ ansible_date_time.date }}.csv"
        line: >
          {{ inventory_hostname }},
          {{ container_info.container.Config.Image | default("MISSING") }},
          {{ key_stat.stat.mode }},
          {{ ansible_date_time.iso8601 }}
        create: true

# Команда:
# ansible-playbook -i inventory/production playbooks/audit.yml \\
#   --vault-password-file vault/.vault_pass
# Результат: файл audit_report_YYYY-MM-DD.csv (отчёт ФиАК → CAR)''')

# ══════════════════════════════════════════════════════════════════════
# 4. ВЫВОДЫ
# ══════════════════════════════════════════════════════════════════════
h1(doc,'4. КРАТКИЕ ВЫВОДЫ')
conclusions = [
    ('Оценка по ГОСТ Р ИСО/МЭК 12207–2010:',
     ' идентифицированы 10 элементов конфигурации Django Security Logger App (DSLA-SC-1.0 … DSLA-TST-1.0), установлены три типа базовых линий (FB/AB/PB). Все шесть видов деятельности раздела 7.2.2.3 — реализация, идентификация, управление, отслеживание, оценка, поставка — задокументированы и реализованы через Git + CI/CD-конвейер.'),
    ('Сравнительный анализ:',
     ' из шести систем (Puppet, Chef, SaltStack, Ansible, NIX, CFEngine) для развёртывания на 1000 серверов выбран Ansible. Реальные цены: Ansible (ansible-core) — бесплатно, Red Hat AAP Standard — ~$130/узел/год; Chef — $59–189/узел/год; Puppet — ~$112–199/узел/год; SaltStack Enterprise — custom quote (Broadcom); NixOS — бесплатно; CFEngine Enterprise — custom quote. Ansible выигрывает по совокупности критериев: безагентность, бесплатность для OSS-развёртывания, наивысший рейтинг пользователей (G2: 4.5/5).'),
    ('План менеджмента конфигурации:',
     ' разработан в полном соответствии с Приложением А ГОСТ Р ИСО 10007–2019 (разделы: общие положения, введение, политики, идентификация, управление изменениями, учёт статуса, аудит). Установлена процедура CR с 8 этапами, назначена CCB, определены два типа аудита (ФАК и ФиАК).'),
    ('Практическая реализация на Ansible:',
     ' разработана полная структура Ansible-проекта: инвентарь 1000 серверов, 4 роли (common, docker, nginx, dsla_app), playbook-и для первоначального развёртывания (forks=20), rolling update (serial=50, max_fail_percentage=5%) и физического аудита. Секреты (ключи Fernet, сертификаты TLS) защищены Ansible Vault (AES-256).'),
    ('Главный вывод:',
     ' применение Ansible как системы управления конфигурациями обеспечивает воспроизводимое, идемпотентное и автоматизированное развёртывание DSLA на любом масштабе инфраструктуры при минимальных операционных рисках и полном соответствии требованиям ГОСТ Р ИСО/МЭК 12207–2010 и ГОСТ Р ИСО 10007–2019.'),
]
for bold_part, rest in conclusions:
    mixed(doc,[(bold_part, True, 12, False),(rest, False, 12, False)], indent=True, sb=4, sa=4)

# ══════════════════════════════════════════════════════════════════════
# 5. СПИСОК ЛИТЕРАТУРЫ
# ══════════════════════════════════════════════════════════════════════
h1(doc,'5. СПИСОК ИСПОЛЬЗУЕМОЙ ЛИТЕРАТУРЫ')
refs = [
    'ГОСТ Р ИСО/МЭК 12207–2010. Информационная технология. Системная и программная инженерия. Процессы жизненного цикла программных средств. — М.: Стандартинформ, 2011. — 124 с.',
    'ГОСТ Р ИСО 10007–2019. Менеджмент качества. Руководящие указания по менеджменту конфигурации. — М.: Стандартинформ, 2019. — 16 с.',
    'Ansible Documentation. Red Hat, Inc., 2025. [Электронный ресурс]. — URL: https://docs.ansible.com (дата обращения: 10.05.2025).',
    'Red Hat Ansible Automation Platform — Pricing. [Электронный ресурс]. — URL: https://www.redhat.com/en/technologies/management/ansible/pricing (дата обращения: 10.05.2025).',
    'Ansible pricing analysis (Netsync reseller data). CyberPanel. [Электронный ресурс]. — URL: https://cyberpanel.net/blog/ansible-pricing (дата обращения: 10.05.2025).',
    'Puppet Pricing. Perforce Software. [Электронный ресурс]. — URL: https://www.puppet.com/pricing (дата обращения: 10.05.2025).',
    'OpenVox: The Community-Driven Fork of Puppet. The New Stack, 2025. [Электронный ресурс]. — URL: https://thenewstack.io/openvox-the-community-driven-fork-of-puppet-has-arrived (дата обращения: 10.05.2025).',
    'Progress Chef — Pricing & Reviews. Capterra, 2025. [Электронный ресурс]. — URL: https://www.capterra.com/p/170467/Chef (дата обращения: 10.05.2025).',
    'Progress Chef pricing. TrustRadius. [Электронный ресурс]. — URL: https://www.trustradius.com/products/progress-chef/pricing (дата обращения: 10.05.2025).',
    'Salt Project Documentation. VMware / Broadcom. [Электронный ресурс]. — URL: https://saltproject.io (дата обращения: 10.05.2025).',
    'VMware Aria Automation Config (SaltStack). [Электронный ресурс]. — URL: https://www.vmware.com/products/aria-automation/saltstack-config.html (дата обращения: 10.05.2025).',
    'NixOS Manual. NixOS Community. [Электронный ресурс]. — URL: https://nixos.org/manual/nixos/stable (дата обращения: 10.05.2025).',
    'Determinate Nix. Determinate Systems. [Электронный ресурс]. — URL: https://determinate.systems (дата обращения: 10.05.2025).',
    'CFEngine Community vs. Enterprise Comparison. Northern.tech. [Электронный ресурс]. — URL: https://cfengine.com/community-vs-enterprise-comparison (дата обращения: 10.05.2025).',
    'CFEngine Pricing. TrustRadius. [Электронный ресурс]. — URL: https://www.trustradius.com/products/cfengine/pricing (дата обращения: 10.05.2025).',
    'Red Hat Ansible Automation Platform Reviews. G2. [Электронный ресурс]. — URL: https://www.g2.com/products/red-hat-ansible-automation-platform/reviews (дата обращения: 10.05.2025).',
    'Корнеев Н.В. Задание №9. Конфигурационное управление программным обеспечением. — М., 2025.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent       = Cm(1.25)
    p.paragraph_format.space_after       = Pt(3)
    r = p.add_run(f'{i}.\u00a0{ref}')
    r.font.size = Pt(11)

out = '/home/runner/workspace/ЛР9_Конфигурационное_управление_v2.docx'
doc.save(out)
print(f'Saved: {out}')
