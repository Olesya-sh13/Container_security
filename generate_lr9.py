from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

# --- Styles helper ---
def add_heading(doc, text, level=1, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_body(doc, text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

# ==================== TITLE PAGE ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('МИНОБРНАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ')
run.font.size = Pt(12)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Федеральное государственное бюджетное образовательное учреждение\nвысшего образования')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
run = p.add_run('Кафедра информационных технологий и кибербезопасности')
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ОТЧЁТ')
run.font.size = Pt(16)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('по лабораторной работе № 9')
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run('«Конфигурационное управление программным обеспечением»')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Дисциплина: Сопровождение программных средств\n')
run.font.size = Pt(12)
run = p.add_run('Вариант: Django Security Logger App\n')
run.font.size = Pt(12)
run = p.add_run('Выполнил: студент группы ___\n')
run.font.size = Pt(12)
run = p.add_run('Проверил: Н.В. Корнеев')
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Москва — 2025')
run.font.size = Pt(12)

doc.add_page_break()

# ==================== 1. ЦЕЛЬ ====================
add_heading(doc, '1. ЦЕЛЬ РАБОТЫ', level=1, size=13)
add_body(doc, 'Получить навыки конфигурационного управления программным обеспечением на примере контейнера для обеспечения безопасности приложения.')

# ==================== 2. ФОРМУЛИРОВКА ЗАДАНИЯ ====================
add_heading(doc, '2. ФОРМУЛИРОВКА ЗАДАНИЯ', level=1, size=13)
add_body(doc, 'Для программного средства по варианту индивидуального задания (Django Security Logger App — система мониторинга и журналирования событий безопасности виртуальных машин) рассмотреть процесс менеджмента конфигурации по ГОСТ Р ИСО/МЭК 12207–2010.')
add_body(doc, 'Руководством компании поставлена задача развернуть программное средство одновременно на 1000 серверов заказчика, а затем автоматически изменять конфигурацию при необходимости. Являясь сотрудником компании, необходимо:')
add_bullet(doc, 'провести оценку процесса менеджмента конфигурации программного средства по ГОСТ Р ИСО/МЭК 12207–2010 (раздел 7.2.2);')
add_bullet(doc, 'провести сравнительный анализ 5–6 систем управления конфигурациями по критериям: функциональность, стоимость, безопасность и отзывы пользователей (Таблица 1);')
add_bullet(doc, 'подготовить план менеджмента конфигурации по ГОСТ Р ИСО 10007–2019;')
add_bullet(doc, 'настроить модель реализации с помощью выбранной системы управления конфигурациями (Ansible).')

# ==================== 3. РЕЗУЛЬТАТЫ ====================
add_heading(doc, '3. РЕЗУЛЬТАТЫ ВЫПОЛНЕНИЯ ЗАДАНИЯ', level=1, size=13)

# --- 3.1 Оценка по ГОСТ 12207 ---
add_heading(doc, '3.1. Оценка процесса менеджмента конфигурации по ГОСТ Р ИСО/МЭК 12207–2010', level=2, size=12)

add_body(doc, 'Согласно ГОСТ Р ИСО/МЭК 12207–2010, процесс менеджмента конфигурации программных средств описан в разделе 7.2.2. Его цель заключается в установлении и сопровождении целостности программных составных частей процесса или проекта и обеспечении их доступности для заинтересованных сторон.')

add_heading(doc, '3.1.1. Идентификация элементов конфигурации Django Security Logger App', level=3, size=12, bold=False)
add_body(doc, 'Для рассматриваемого программного средства были определены следующие элементы конфигурации (ЭК):')
items_ci = [
    ('ЭК-01', 'Исходный код приложения (logger_app/, django_project/)'),
    ('ЭК-02', 'Файлы шаблонов HTML (templates/logger_app/)'),
    ('ЭК-03', 'Модуль шифрования журналов (logger_app/encryption.py, logger_app/data/secret.key)'),
    ('ЭК-04', 'Модуль моделирования событий безопасности (logger_app/data/utils.py)'),
    ('ЭК-05', 'База данных (db.sqlite3)'),
    ('ЭК-06', 'Политика доступа (logger_app/migrations/policy.json)'),
    ('ЭК-07', 'Конфигурация проекта Django (django_project/settings.py)'),
    ('ЭК-08', 'Docker-инфраструктура (Dockerfile, docker-compose.yml, nginx/nginx.conf)'),
    ('ЭК-09', 'Агент журналирования КЖД v2.0 (docker/logger/)'),
    ('ЭК-10', 'Тесты (tests/)'),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
set_cell_text(hdr[0], 'Идентификатор', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(hdr[1], 'Наименование ЭК', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(hdr[2], 'Версия / Базовая линия', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
for c in hdr: set_cell_bg(c, 'D9E1F2')
for id_, name in items_ci:
    row = tbl.add_row().cells
    set_cell_text(row[0], id_, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row[1], name, size=10)
    set_cell_text(row[2], 'v1.0 / baseline-2025', size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_heading(doc, '3.1.2. Виды деятельности процесса менеджмента конфигурации (ГОСТ Р ИСО/МЭК 12207–2010, п. 7.2.2.3)', level=3, size=12, bold=False)

activities = [
    ('7.2.2.3.1 Реализация процесса',
     'Разработан план менеджмента конфигурации (ПМК) Django Security Logger App. В плане описаны: состав ЭК, стратегия управления конфигурацией с применением системы Ansible, порядок внесения изменений, ответственные исполнители. ПМК утверждён руководителем проекта и хранится в репозитории Git.'),
    ('7.2.2.3.2 Идентификация конфигурации',
     'Для каждого ЭК присвоен уникальный идентификатор (ЭК-01 … ЭК-10). Версии отслеживаются через систему контроля версий Git (теги v1.0, v1.1 и т.д.). Базовые линии фиксируются после каждого релиза. Схема идентификации: <проект>-<компонент>-<версия> (например, DSLA-SCR-1.0).'),
    ('7.2.2.3.3 Управление конфигурацией',
     'Все изменения проходят через процедуру заявки на изменение (Change Request, CR). Заявки регистрируются в системе GitHub Issues, оцениваются ответственным исполнителем, реализуются в отдельной ветке Git, проходят code review и автоматическое тестирование (pytest), после чего сливаются в основную ветку. Критические изменения, влияющие на безопасность (модуль шифрования, политика доступа), требуют дополнительного аудита.'),
    ('7.2.2.3.4 Отслеживание состояния конфигурации',
     'Статус каждого ЭК регистрируется в журнале состояния конфигурации (Configuration Status Log, CSL). Отчёты формируются ежеквартально и при каждом выпуске релиза. Отчёты содержат: список активных ЭК, их текущие версии, историю изменений, статус открытых CR.'),
    ('7.2.2.3.5 Оценка конфигурации',
     'Функциональная завершённость ЭК проверяется с помощью набора автоматизированных тестов (pytest; 41 тест-кейс в tests/test_https_container.py, тесты шифрования, тесты моделей и представлений). Физическая завершённость проверяется при аудите конфигурации путём сравнения развёрнутых версий с эталоном в репозитории.'),
    ('7.2.2.3.6 Поставка и менеджмент выпуска',
     'Выпуски программного средства формируются в виде Docker-образов и публикуются в Container Registry. Каждый выпуск сопровождается тегом Git, CHANGELOG и подписанным манифестом образа. Для массового развёртывания на 1000 серверов используется система Ansible с централизованным хранилищем playbook-ов.'),
]
for title, text in activities:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(title + '. ')
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(text)
    run2.font.size = Pt(12)

add_body(doc, 'Таким образом, процесс менеджмента конфигурации Django Security Logger App полностью соответствует требованиям ГОСТ Р ИСО/МЭК 12207–2010. Все шесть видов деятельности (реализация, идентификация, управление, отслеживание, оценка, поставка) реализованы и документально подтверждены.')

# --- 3.2 Сравнительный анализ ---
add_heading(doc, '3.2. Сравнительный анализ систем управления конфигурациями', level=2, size=12)

add_body(doc, 'Для решения задачи развёртывания программного средства на 1000 серверов и последующего автоматического управления конфигурацией был проведён сравнительный анализ шести систем управления конфигурациями по четырём критериям.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Таблица 1 — Сравнительный анализ систем управления конфигурациями')
run.bold = True
run.font.size = Pt(11)

# Table 1
headers = ['№', 'Система', 'Функциональность', 'Стоимость', 'Безопасность', 'Отзывы пользователей']
col_widths = [Cm(0.8), Cm(2.2), Cm(4.0), Cm(3.0), Cm(3.5), Cm(4.2)]

systems = [
    ['1', 'Puppet',
     'Декларативный язык DSL (Puppet DSL). Управление агентами (Puppet Agents) на клиентах. Поддержка модульной архитектуры (Puppet Forge — >6000 модулей). Отчёты о состоянии узлов, каталоги конфигурации.',
     'Puppet Community — бесплатно. Puppet Enterprise — от $120/узел/год. Puppet DRA — корпоративная лицензия.',
     'Централизованный Puppet Master. Взаимная аутентификация по TLS/SSL (PKI). RBAC в Enterprise. Подписание сертификатов агентов. Аудит изменений.',
     'Декларативный подход позволяет описывать желаемое состояние системы. Сложная кривая обучения DSL. Высокая зрелость — используется в крупных enterprise-проектах. [1]'],
    ['2', 'Chef',
     'Императивный язык Ruby (Recipes/Cookbooks). Поддержка сложных условных сценариев. Chef Supermarket — библиотека кулинарных книг. Test Kitchen для тестирования.',
     'Chef Infra (Community) — бесплатно (Apache 2.0). Chef Infra Server — бесплатно до 25 узлов, далее коммерческая. Progress Chef — enterprise от $137/узел/год.',
     'Шифрование данных атрибутов (Chef Vault). Аутентификация Chef Server по SSL. RBAC. Подпись запросов API.',
     'Мощный инструмент для разработчиков, знакомых с Ruby. Гибкость приводит к сложности. Часто применяется в DevOps с CI/CD. Хорошая интеграция с облаками AWS/Azure/GCP. [2]'],
    ['3', 'SaltStack\n(Salt)',
     'Язык YAML (States/Pillars) + Jinja2. Быстрая передача команд через ZeroMQ (push-модель). Событийная система (Salt Reactor). Поддержка агентного (minion) и безагентного (SSH) режима.',
     'Salt Open Source — бесплатно (Apache 2.0). VMware Aria Automation Config (ex-SaltStack Enterprise) — коммерческая лицензия по запросу.',
     'Шифрование AES для передачи данных. Аутентификация Salt Master/Minion по ключам. Pillar для хранения секретов. Безагентный SSH-режим без открытых портов.',
     'Высокая производительность при работе с тысячами узлов. YAML прост для понимания. Реактивная событийная система — уникальное преимущество. Сложность настройки при большом количестве minion-ов. [3]'],
    ['4', 'Ansible',
     'Декларативный YAML (Playbooks/Roles). Безагентная архитектура (SSH/WinRM). Ansible Galaxy — >50000 ролей. Идемпотентность модулей. AWX/Ansible Tower — веб-интерфейс и API.',
     'Ansible Community — полностью бесплатно (GPL). Red Hat Ansible Automation Platform — от $14000/год (50 управляемых узлов). AWX — бесплатный upstream.',
     'Передача данных по SSH (без агентов на узлах). Ansible Vault — шифрование секретов (AES-256). RBAC в Ansible Tower. Минимальная поверхность атаки (нет демона на клиентах).',
     'Самая популярная система управления конфигурациями (>26 млн загрузок). Простота — playbook читается как документация. Огромное сообщество. Может замедляться на >1000 узлов без Ansible Tower. [4]'],
    ['5', 'NIX /\nNixOS',
     'Функциональный язык Nix. Полностью воспроизводимые сборки (reproducible builds). Атомарные обновления и откаты. NixOS Modules — декларативная конфигурация ОС. Flakes для управления зависимостями.',
     'NixOS — полностью бесплатно (LGPL/MIT). Коммерческой версии нет. Поддержка через Determinate Systems (коммерческая).',
     'Изоляция пакетов в /nix/store (хэш-адресация). Неизменяемая инфраструктура. Отсутствие глобальных зависимостей снижает риск атак на цепочку поставок. Воспроизводимость сборок — верифицируемый результат.',
     'Высокий порог вхождения: язык Nix нетривиален. Идеален для воспроизводимых окружений разработки и облачных платформ. Растущее сообщество. Применяется в исследовательских проектах и облачных окружениях (AWS, GCP). [5]'],
    ['6', 'CFEngine',
     'Декларативный язык CFEngine DSL (Promise Theory). Один из старейших инструментов (с 1993). Легковесный агент. Автономный режим работы (агент действует независимо от центра при потере связи).',
     'CFEngine Community — бесплатно (GPL). CFEngine Enterprise — коммерческая (по запросу, ~$50/узел/год).',
     'PKI-аутентификация. Минимальный сетевой трафик (шифрование TLS). Автономность агентов — нет единой точки отказа. Поддержка FIPS 140-2.',
     'Применяется в высоконагруженных средах с тысячами серверов (в т.ч. финансовый сектор). Крутая кривая обучения. Устаревшая документация. Медленно развивается. Высокая надёжность в автономном режиме. [6]'],
]

tbl2 = doc.add_table(rows=1, cols=6)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for i, w in enumerate(col_widths):
    for cell in tbl2.columns[i].cells:
        cell.width = w

hdr2 = tbl2.rows[0].cells
for i, h in enumerate(headers):
    set_cell_text(hdr2[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(hdr2[i], '4472C4')
    hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

alt_color = 'DEEAF1'
for idx, row_data in enumerate(systems):
    row = tbl2.add_row().cells
    bg = alt_color if idx % 2 == 0 else 'FFFFFF'
    for i, cell_text in enumerate(row_data):
        set_cell_text(row[i], cell_text, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_cell_bg(row[i], bg)

# Sources
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run('Источники к Таблице 1:')
run.bold = True
run.font.size = Pt(10)

sources = [
    '[1] Puppet vs Ansible: https://kurshub.ru/journal/blog/ansible-ili-puppet-razbor-silnyh-i-slabyh-storon/',
    '[2] Chef Documentation: https://docs.chef.io/',
    '[3] SaltStack Overview: https://docs.saltproject.io/en/latest/topics/about_salt_project/index.html',
    '[4] Ansible Documentation: https://docs.ansible.com/',
    '[5] NixOS Manual: https://nixos.org/manual/nixos/stable/',
    '[6] CFEngine Documentation: https://docs.cfengine.com/',
]
for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(s)
    run.font.size = Pt(9)

add_body(doc, 'По результатам сравнительного анализа для решения задачи развёртывания Django Security Logger App на 1000 серверов выбрана система Ansible. Основные преимущества: безагентная архитектура (не требует установки ПО на серверах), простой YAML-синтаксис playbook, встроенное шифрование секретов (Ansible Vault), широкая поддержка Linux/Docker-окружений, бесплатная open-source версия. При необходимости масштабируется с Ansible Tower/AWX.')

# --- 3.3 План менеджмента конфигурации ---
add_heading(doc, '3.3. План менеджмента конфигурации по ГОСТ Р ИСО 10007–2019', level=2, size=12)

# 3.3.1 Общие положения
add_heading(doc, '3.3.1. Общие положения', level=3, size=12, bold=True)
add_body(doc, 'Настоящий план менеджмента конфигурации (ПМК) распространяется на программное средство Django Security Logger App (DSLA) — веб-приложение для мониторинга событий безопасности виртуальных машин с ролевым разграничением доступа, шифрованием журналов и аналитической панелью.')
add_body(doc, 'ПМК разработан в соответствии с ГОСТ Р ИСО 10007–2019 «Менеджмент качества. Руководящие указания по менеджменту конфигурации» и ГОСТ Р ИСО/МЭК 12207–2010 «Процессы жизненного цикла программных средств» (п. 7.2.2).')
add_body(doc, 'Сфера применения: разработка, тестирование, развёртывание и сопровождение DSLA на 1000 серверов заказчика. Стадии жизненного цикла: концепция → разработка → тестирование → развёртывание → сопровождение.')

# 3.3.2 Введение
add_heading(doc, '3.3.2. Введение', level=3, size=12, bold=True)
add_body(doc, 'Целью ПМК является обеспечение целостности, прослеживаемости и управляемости всех элементов конфигурации DSLA на всём жизненном цикле продукта.')
add_body(doc, 'Объекты управления: программный код, конфигурационные файлы, Docker-образы, политики доступа, ключи шифрования, документация, тестовые наборы.')
add_body(doc, 'Инструменты менеджмента конфигурации:')
add_bullet(doc, 'Git (GitHub) — система контроля версий исходного кода;')
add_bullet(doc, 'Ansible — система управления конфигурациями целевых серверов;')
add_bullet(doc, 'Ansible Vault — шифрование секретных данных в playbook;')
add_bullet(doc, 'Docker / Container Registry — упаковка и распространение релизов;')
add_bullet(doc, 'GitHub Actions — CI/CD-конвейер сборки, тестирования, публикации образов;')
add_bullet(doc, 'GitHub Issues — регистрация и отслеживание заявок на изменения.')
add_body(doc, 'График ключевых видов деятельности:')
add_bullet(doc, 'Идентификация ЭК и установка базовых линий — при старте проекта и при каждом релизе;')
add_bullet(doc, 'Регистрация изменений — непрерывно (по мере поступления CR);')
add_bullet(doc, 'Отчёты о состоянии конфигурации — ежеквартально;')
add_bullet(doc, 'Аудит конфигурации — перед каждым производственным релизом и не реже 1 раза в год.')

# 3.3.3 Политики
add_heading(doc, '3.3.3. Политики', level=3, size=12, bold=True)
add_body(doc, 'Политика в области менеджмента конфигурации DSLA включает следующие принципы:')
add_bullet(doc, 'Все ЭК находятся под контролем версий Git. Прямые коммиты в ветку main запрещены — изменения вносятся только через Pull Request с обязательным code review.')
add_bullet(doc, 'Каждый релиз сопровождается тегом Git формата vMAJOR.MINOR.PATCH и Docker-образом с тем же тегом.')
add_bullet(doc, 'Секретные данные (ключи шифрования, пароли БД, сертификаты) хранятся в Ansible Vault и никогда не попадают в открытый репозиторий.')
add_bullet(doc, 'Ответственность за одобрение изменений в критических ЭК (модуль шифрования, политика доступа) несёт менеджер безопасности.')
add_bullet(doc, 'Квалификация: все участники процесса прошли обучение по Git Flow, Ansible и политике безопасности компании.')
add_bullet(doc, 'Терминология соответствует ГОСТ Р ИСО 10007–2019 и ГОСТ Р ИСО/МЭК 12207–2010.')

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.25)
run = p.add_run('Распределение ответственности:')
run.bold = True
run.font.size = Pt(12)

roles = [
    ('Менеджер конфигурации', 'Разработка и сопровождение ПМК, утверждение ЭК, координация процесса'),
    ('Ответственный исполнитель (Configuration Control Board, CCB)', 'Оценка и одобрение заявок на изменение'),
    ('Разработчики', 'Соблюдение процедур ветвления Git, документирование изменений'),
    ('DevOps-инженер', 'Сопровождение Ansible playbook, CI/CD-конвейеров, Docker Registry'),
    ('Менеджер безопасности', 'Аудит критических ЭК, управление ключами Ansible Vault'),
    ('Тестировщик', 'Верификация корректности изменений (pytest, интеграционные тесты)'),
]
tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = 'Table Grid'
set_cell_text(tbl3.rows[0].cells[0], 'Роль', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(tbl3.rows[0].cells[1], 'Ответственность', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
for c in tbl3.rows[0].cells: set_cell_bg(c, 'D9E1F2')
for role, resp in roles:
    r = tbl3.add_row().cells
    set_cell_text(r[0], role, size=10)
    set_cell_text(r[1], resp, size=10)
doc.add_paragraph()

# 3.3.4 Идентификация конфигурации
add_heading(doc, '3.3.4. Идентификация конфигурации', level=3, size=12, bold=True)
add_body(doc, 'Схема идентификации ЭК: DSLA-<КОД>-<ВЕРСИЯ>, где КОД — двухбуквенный код компонента, ВЕРСИЯ — номер в формате MAJOR.MINOR.PATCH.')
add_body(doc, 'Пример: DSLA-SC-1.2.0 — исходный код приложения, версия 1.2.0.')
add_body(doc, 'Базовые конфигурации:')
add_bullet(doc, 'Functional Baseline (FB) — устанавливается при завершении технического задания (требования к DSLA);')
add_bullet(doc, 'Allocated Baseline (AB) — устанавливается при завершении проектирования архитектуры;')
add_bullet(doc, 'Product Baseline (PB) — устанавливается при каждом производственном релизе.')
add_body(doc, 'Все базовые конфигурации фиксируются тегами Git и записываются в реестр ЭК (Configuration Item Register, CIR). Статус пересмотра обозначается суффиксом: -alpha, -beta, -rc.N, -release.')

# 3.3.5 Управление изменениями
add_heading(doc, '3.3.5. Управление изменениями', level=3, size=12, bold=True)
add_body(doc, 'Процедура управления изменениями включает следующие этапы:')

steps_change = [
    ('1. Инициирование', 'Заинтересованная сторона (разработчик, заказчик, служба безопасности) создаёт заявку на изменение (CR) в GitHub Issues с описанием проблемы, обоснованием и категорией изменения (корректирующее / улучшающее / адаптивное).'),
    ('2. Идентификация и документирование', 'CR присваивается уникальный номер. Указываются затронутые ЭК, текущая версия, описание предполагаемого изменения и потенциальные риски.'),
    ('3. Оценка (CCB)', 'Ответственный исполнитель (CCB) оценивает технические преимущества, риски, влияние на безопасность, трудозатраты и сроки. Решение фиксируется в CR.'),
    ('4. Одобрение / отклонение', 'CCB принимает решение: одобрить, отклонить или отложить. Уведомление направляется всем заинтересованным сторонам.'),
    ('5. Реализация', 'Разработчик создаёт ветку Git формата feature/CR-NNN, вносит изменения, покрывает их тестами (pytest), открывает Pull Request.'),
    ('6. Верификация', 'CI/CD-конвейер (GitHub Actions) автоматически выполняет сборку и тесты. После прохождения тестов проводится code review.'),
    ('7. Внедрение', 'После одобрения PR изменения сливаются в main, формируется новый релиз (тег Git + Docker-образ), Ansible playbook автоматически обновляет целевые серверы.'),
    ('8. Закрытие CR', 'CR закрывается с указанием результата внедрения и обновлённой версии ЭК. Документация обновляется.'),
]
for step, desc in steps_change:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(step + ': ')
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(desc)
    run2.font.size = Pt(12)

add_body(doc, 'Категории изменений: Minor (не влияет на интерфейс и совместимость), Major (изменение интерфейса или поведения), Critical (влияет на безопасность или целостность данных). Critical-изменения требуют дополнительного аудита безопасности.')

# 3.3.6 Учёт статуса конфигурации
add_heading(doc, '3.3.6. Учёт статуса конфигурации', level=3, size=12, bold=True)
add_body(doc, 'Учёт статуса конфигурации обеспечивает прослеживаемость всех изменений на всех стадиях жизненного цикла DSLA. Документированная информация включает:')
add_bullet(doc, 'Реестр ЭК (CIR) — список всех ЭК с текущей версией, статусом и историей изменений;')
add_bullet(doc, 'Журнал CR — все поданные заявки на изменение с их статусом (открыта / в работе / одобрена / внедрена / закрыта);')
add_bullet(doc, 'Журнал выпусков (Release Log) — перечень выпущенных версий с датой, составом изменений и ответственным;')
add_bullet(doc, 'Отчёты о состоянии конфигурации (CSR) — ежеквартальные сводные отчёты.')
add_body(doc, 'Данные хранятся в защищённом Git-репозитории с ограниченным доступом. Резервное копирование — ежедневно. Целостность данных обеспечивается криптографическими подписями коммитов (GPG).')
add_body(doc, 'Отчёт о состоянии конфигурации (CSR) содержит: перечень активных ЭК с версиями; число CR за период (открытых / закрытых); список последних выпусков; список ЭК, требующих обновления; результаты последнего аудита.')

# 3.3.7 Аудит конфигурации
add_heading(doc, '3.3.7. Аудит конфигурации', level=3, size=12, bold=True)
add_body(doc, 'В соответствии с ГОСТ Р ИСО 10007–2019 (п. 5.5.4) проводятся два типа аудита:')
add_bullet(doc, 'Функциональный аудит конфигурации (ФАК): верификация того, что ЭК достиг функциональных и рабочих характеристик, указанных в ТЗ. Инструмент — автоматизированное тестирование (pytest). Периодичность: перед каждым производственным релизом.')
add_bullet(doc, 'Физический аудит конфигурации (ФиАК): верификация соответствия развёрнутых артефактов эталонным образам в Container Registry. Инструмент — сравнение хэшей Docker-образов, Ansible ad-hoc команды. Периодичность: не реже 1 раза в год и после инцидентов безопасности.')
add_body(doc, 'Ответственный за аудит: менеджер конфигурации совместно с менеджером безопасности.')
add_body(doc, 'Результаты аудита оформляются в виде Отчёта об аудите конфигурации (CAR) и направляются руководству проекта. При выявлении несоответствий инициируется CR категории Critical.')

doc.add_page_break()

# --- 3.4 Настройка модели реализации с Ansible ---
add_heading(doc, '3.4. Настройка модели реализации с помощью Ansible', level=2, size=12)
add_heading(doc, 'Приложение: практический пример развёртывания Django Security Logger App на 1000 серверов с помощью Ansible', level=3, size=12, bold=False)

add_body(doc, 'Ansible выбран как инструмент управления конфигурацией для Django Security Logger App на основании результатов сравнительного анализа (Таблица 1). Ниже представлена полная модель реализации.')

add_heading(doc, 'Структура проекта Ansible', level=3, size=12, bold=True)
add_body(doc, 'Рекомендуемая структура директорий Ansible-проекта:', first_indent=True)

struct = '''dsla-ansible/
├── inventory/
│   ├── production/
│   │   ├── hosts.ini          # 1000 серверов заказчика
│   │   └── group_vars/
│   │       ├── all.yml        # общие переменные
│   │       └── all_vault.yml  # зашифрованные секреты (Ansible Vault)
│   └── staging/
│       └── hosts.ini          # тестовые серверы
├── roles/
│   ├── common/                # базовая настройка ОС
│   ├── docker/                # установка Docker + Docker Compose
│   ├── dsla_app/              # развёртывание Django Security Logger App
│   └── nginx/                 # настройка Nginx + TLS
├── playbooks/
│   ├── site.yml               # основной playbook (полное развёртывание)
│   ├── update.yml             # обновление конфигурации без остановки сервиса
│   └── audit.yml              # аудит состояния конфигурации
├── ansible.cfg                # конфигурация Ansible
└── requirements.yml           # зависимости ролей (Ansible Galaxy)'''

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(6)
run = p.add_run(struct)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Файл инвентаря (inventory/production/hosts.ini)', level=3, size=12, bold=True)
inv = '''[dsla_servers]
server[001:500].prod.example.com  ansible_user=deploy  ansible_port=22
server[501:1000].prod.example.com ansible_user=deploy  ansible_port=22

[dsla_servers:vars]
ansible_python_interpreter=/usr/bin/python3
ansible_ssh_private_key_file=~/.ssh/dsla_deploy_key'''

p = doc.add_paragraph()
run = p.add_run(inv)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Зашифрованные переменные (group_vars/all_vault.yml)', level=3, size=12, bold=True)
add_body(doc, 'Файл хранится в зашифрованном виде (Ansible Vault, AES-256). Создаётся командой: ansible-vault create inventory/production/group_vars/all_vault.yml', first_indent=False)
vault_ex = '''# Содержимое (после расшифровки):
vault_django_secret_key: "50_символов_случайного_ключа"
vault_fernet_key: "ключ_шифрования_журналов_Fernet_base64"
vault_db_password: "пароль_БД"
vault_ssl_cert_path: "/etc/nginx/certs/dsla.crt"
vault_ssl_key_path: "/etc/nginx/certs/dsla.key"'''

p = doc.add_paragraph()
run = p.add_run(vault_ex)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Основные переменные (group_vars/all.yml)', level=3, size=12, bold=True)
vars_ex = '''dsla_version: "1.2.0"
dsla_image: "registry.example.com/dsla:{{ dsla_version }}"
dsla_port: 5000
nginx_port_http: 80
nginx_port_https: 443
dsla_log_dir: "/var/log/dsla"
dsla_data_dir: "/opt/dsla/data"
dsla_replicas: 1'''

p = doc.add_paragraph()
run = p.add_run(vars_ex)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Основной playbook (playbooks/site.yml)', level=3, size=12, bold=True)
site_yml = '''---
- name: Deploy Django Security Logger App
  hosts: dsla_servers
  become: true
  vars_files:
    - "{{ inventory_dir }}/group_vars/all_vault.yml"
  
  roles:
    - role: common        # Обновление ОС, настройка часового пояса, fail2ban
    - role: docker        # Установка Docker CE, docker-compose plugin
    - role: nginx         # Установка Nginx, копирование TLS-сертификатов
    - role: dsla_app      # Развёртывание Django Security Logger App'''

p = doc.add_paragraph()
run = p.add_run(site_yml)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Роль dsla_app: основные задачи (roles/dsla_app/tasks/main.yml)', level=3, size=12, bold=True)
tasks_yml = '''---
- name: Создать директории приложения
  file:
    path: "{{ item }}"
    state: directory
    owner: deploy
    mode: '0750'
  loop:
    - "{{ dsla_data_dir }}"
    - "{{ dsla_log_dir }}"
    - "{{ dsla_data_dir }}/certs"

- name: Скопировать ключ Fernet из Vault
  copy:
    content: "{{ vault_fernet_key }}"
    dest: "{{ dsla_data_dir }}/secret.key"
    owner: deploy
    mode: '0600'

- name: Скопировать политику доступа
  template:
    src: policy.json.j2
    dest: "{{ dsla_data_dir }}/policy.json"
    owner: deploy
    mode: '0640'

- name: Авторизоваться в Container Registry
  docker_login:
    registry: registry.example.com
    username: "{{ vault_registry_user }}"
    password: "{{ vault_registry_password }}"

- name: Развернуть Django Security Logger App (docker-compose)
  docker_compose:
    project_name: dsla
    definition:
      version: '3.8'
      services:
        app:
          image: "{{ dsla_image }}"
          restart: unless-stopped
          environment:
            SECRET_KEY: "{{ vault_django_secret_key }}"
            DJANGO_SETTINGS_MODULE: "django_project.settings"
          volumes:
            - "{{ dsla_data_dir }}:/app/logger_app/data"
            - "{{ dsla_log_dir }}:/app/logger_app/logs"
          ports:
            - "127.0.0.1:{{ dsla_port }}:5000"
    state: present
    pull: true
  notify: Restart DSLA

- name: Выполнить миграции Django
  docker_container_exec:
    container: dsla_app_1
    command: python manage.py migrate --noinput
  when: dsla_run_migrations | default(true)

- name: Проверить доступность приложения
  uri:
    url: "http://127.0.0.1:{{ dsla_port }}/logger/"
    status_code: 200
    timeout: 30
  retries: 5
  delay: 10'''

p = doc.add_paragraph()
run = p.add_run(tasks_yml)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Playbook обновления конфигурации (playbooks/update.yml)', level=3, size=12, bold=True)
add_body(doc, 'Для автоматического изменения конфигурации при необходимости (в соответствии с задачей компании):')
update_yml = '''---
- name: Обновить конфигурацию DSLA без остановки сервиса
  hosts: dsla_servers
  become: true
  serial: 50          # Обновлять по 50 серверов за раз (rolling update)
  max_fail_percentage: 5   # Остановить, если >5% серверов упали
  vars_files:
    - "{{ inventory_dir }}/group_vars/all_vault.yml"

  tasks:
    - name: Обновить образ приложения
      docker_image:
        name: "{{ dsla_image }}"
        source: pull
        force_source: true

    - name: Перезапустить контейнер с новым образом
      docker_compose:
        project_name: dsla
        restarted: true
        pull: true

    - name: Проверить здоровье приложения после обновления
      uri:
        url: "http://127.0.0.1:{{ dsla_port }}/logger/"
        status_code: 200
      retries: 5
      delay: 15'''

p = doc.add_paragraph()
run = p.add_run(update_yml)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Playbook аудита конфигурации (playbooks/audit.yml)', level=3, size=12, bold=True)
audit_yml = '''---
- name: Аудит физической конфигурации DSLA
  hosts: dsla_servers
  become: true
  gather_facts: true

  tasks:
    - name: Получить информацию о запущенном образе
      docker_container_info:
        name: dsla_app_1
      register: container_info

    - name: Проверить соответствие версии эталону
      assert:
        that:
          - container_info.container.Config.Image == dsla_image
        fail_msg: "НЕСООТВЕТСТВИЕ: сервер {{ inventory_hostname }} запускает {{ container_info.container.Config.Image }}"
        success_msg: "OK: версия соответствует {{ dsla_image }}"

    - name: Сохранить результаты аудита
      local_action:
        module: lineinfile
        path: "audit_report_{{ ansible_date_time.date }}.csv"
        line: "{{ inventory_hostname }},{{ container_info.container.Config.Image }},{{ ansible_date_time.iso8601 }}"
        create: true'''

p = doc.add_paragraph()
run = p.add_run(audit_yml)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Команды запуска', level=3, size=12, bold=True)
cmds = '''# Полное развёртывание на 1000 серверов:
ansible-playbook -i inventory/production playbooks/site.yml \
  --vault-password-file ~/.vault_pass --forks 20

# Обновление конфигурации (rolling update):
ansible-playbook -i inventory/production playbooks/update.yml \
  -e "dsla_version=1.3.0" --vault-password-file ~/.vault_pass

# Аудит физической конфигурации:
ansible-playbook -i inventory/production playbooks/audit.yml \
  --vault-password-file ~/.vault_pass

# Проверка связи с узлами:
ansible dsla_servers -i inventory/production -m ping'''

p = doc.add_paragraph()
run = p.add_run(cmds)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_body(doc, 'Параметр --forks 20 означает одновременное подключение к 20 серверам. Для 1000 серверов полное развёртывание займёт порядка 50 последовательных «волн». Rolling update с serial: 50 обеспечивает постепенное обновление без полной остановки сервиса.')

# ==================== 4. ВЫВОДЫ ====================
add_heading(doc, '4. КРАТКИЕ ВЫВОДЫ', level=1, size=13)

add_body(doc, 'В ходе выполнения лабораторной работы были получены следующие результаты:')

conclusions = [
    'Проведена оценка процесса менеджмента конфигурации Django Security Logger App по ГОСТ Р ИСО/МЭК 12207–2010 (раздел 7.2.2). Идентифицированы 10 элементов конфигурации, описаны все шесть видов деятельности: реализация процесса, идентификация конфигурации, управление конфигурацией, отслеживание состояния, оценка и поставка/управление выпуском.',
    'Проведён сравнительный анализ шести систем управления конфигурациями: Puppet, Chef, SaltStack, Ansible, NIX/NixOS, CFEngine. По критериям функциональности, стоимости, безопасности и отзывов пользователей для задачи массового развёртывания выбрана система Ansible благодаря безагентной архитектуре, простоте YAML-синтаксиса, встроенному шифрованию Ansible Vault и широкой экосистеме ролей.',
    'Разработан план менеджмента конфигурации по ГОСТ Р ИСО 10007–2019, включающий разделы: общие положения, введение, политики, идентификация конфигурации, управление изменениями, учёт статуса конфигурации и аудит конфигурации.',
    'Настроена модель реализации с помощью Ansible: разработана структура проекта, инвентарь из 1000 серверов, роли (common, docker, nginx, dsla_app), playbook-и для первоначального развёртывания, rolling update и аудита конфигурации. Секретные данные защищены Ansible Vault (AES-256).',
    'Применение системы управления конфигурациями Ansible обеспечивает воспроизводимость, идемпотентность и автоматизацию процесса развёртывания и обновления Django Security Logger App на инфраструктуре любого масштаба при минимальных операционных рисках.',
]
for i, c in enumerate(conclusions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{i}. {c}')
    run.font.size = Pt(12)

# ==================== 5. СПИСОК ЛИТЕРАТУРЫ ====================
add_heading(doc, '5. СПИСОК ИСПОЛЬЗУЕМОЙ ЛИТЕРАТУРЫ', level=1, size=13)

refs = [
    'ГОСТ Р ИСО/МЭК 12207–2010. Информационная технология. Системная и программная инженерия. Процессы жизненного цикла программных средств. — М.: Стандартинформ, 2011.',
    'ГОСТ Р ИСО 10007–2019. Менеджмент качества. Руководящие указания по менеджменту конфигурации. — М.: Стандартинформ, 2019.',
    'Ansible Documentation. Red Hat, Inc. [Электронный ресурс]. — Режим доступа: https://docs.ansible.com/ (дата обращения: 10.05.2025).',
    'Puppet Documentation. Puppet, Inc. [Электронный ресурс]. — Режим доступа: https://www.puppet.com/docs (дата обращения: 10.05.2025).',
    'Chef Infra Documentation. Progress Software. [Электронный ресурс]. — Режим доступа: https://docs.chef.io/ (дата обращения: 10.05.2025).',
    'SaltStack (Salt) Documentation. VMware, Inc. [Электронный ресурс]. — Режим доступа: https://docs.saltproject.io/ (дата обращения: 10.05.2025).',
    'NixOS Manual. NixOS Community. [Электронный ресурс]. — Режим доступа: https://nixos.org/manual/nixos/stable/ (дата обращения: 10.05.2025).',
    'CFEngine Documentation. Northern.tech. [Электронный ресурс]. — Режим доступа: https://docs.cfengine.com/ (дата обращения: 10.05.2025).',
    'Сравнительный анализ Ansible и Puppet. [Электронный ресурс]. — Режим доступа: https://kurshub.ru/journal/blog/ansible-ili-puppet-razbor-silnyh-i-slabyh-storon/ (дата обращения: 10.05.2025).',
    'Корнеев Н.В. Задание №9. Конфигурационное управление программным обеспечением. — М., 2025.',
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'{i}. {r}')
    run.font.size = Pt(12)

doc.save('/home/runner/workspace/ЛР9_Конфигурационное_управление.docx')
print('Document saved successfully.')
